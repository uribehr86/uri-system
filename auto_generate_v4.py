import os
import openpyxl
from docx import Document

def generate():
    desktop = os.path.join(os.environ['USERPROFILE'], 'Desktop')
    template_path = os.path.join(desktop, 'forms load.docx')
    excel_path = os.path.join(desktop, 'uri test1.xlsx')
    output_path = os.path.join(desktop, 'FINAL_FORMS.docx')

    if not os.path.exists(template_path) or not os.path.exists(excel_path):
        print("Missing files!")
        return

    wb = openpyxl.load_workbook(excel_path)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    headers = rows[0]
    data_rows = rows[1:]
    col_map = {str(h).strip(): i for i, h in enumerate(headers) if h}

    # We will build the final doc by starting with a blank one and 
    # instead of element-level copy, we will just use a simpler replacement if possible.
    # Actually, the most reliable way without external libs is to 
    # use the first student as the base and then try to append.
    
    final_doc = None

    for idx, row in enumerate(data_rows):
        if not any(row): continue
        
        # Load a fresh template for each student
        student_doc = Document(template_path)
        
        def val(header):
            i = col_map.get(header)
            return str(row[i]).strip() if i is not None and row[i] is not None and str(row[i]) != 'None' else ""

        name = val('שם נבחן')
        tz = val('תעודת זהות')
        user = val('קוד משתמש')
        pwd = val('סיסמה')
        notes = val('התאמות')

        # Update Table 0
        if len(student_doc.tables) > 0:
            t = student_doc.tables[0]
            for r in t.rows:
                if len(r.cells) < 2: continue
                txt = r.cells[1].text
                if 'שם נבחן' in txt: r.cells[0].text = name
                elif 'תעודת זהות' in txt: r.cells[0].text = tz
                elif 'קוד משתמש' in txt: r.cells[0].text = user
                elif 'סיסמה' in txt: r.cells[0].text = pwd
                elif 'התאמות' in txt: r.cells[0].text = notes

        # Update Table 1
        if len(student_doc.tables) > 1:
            t = student_doc.tables[1]
            if len(t.rows) >= 2:
                t.cell(1, 0).text = name
                t.cell(1, 1).text = tz

        if final_doc is None:
            final_doc = student_doc
        else:
            final_doc.add_page_break()
            # Append paragraphs and tables
            for p in student_doc.paragraphs:
                new_p = final_doc.add_paragraph(p.text, p.style)
                new_p.alignment = p.alignment
            for t in student_doc.tables:
                # This is a bit limited but safer than element copy
                # For this specific template, we just want the content
                pass 
            # Wait, manual append is hard. Let's try the element copy again but BETTER.
            for element in student_doc.element.body:
                if element.tag.endswith('sectPr'): continue
                final_doc.element.body.append(element)

    if final_doc:
        final_doc.save(output_path)
        print(f"Generated {output_path}")

if __name__ == "__main__":
    generate()
