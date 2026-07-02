import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace', write_through=True)
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace', write_through=True)

from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file, jsonify
from werkzeug.security import generate_password_hash, check_password_hash
import psycopg2
import sqlite3

from psycopg2 import pool
from psycopg2.extras import RealDictCursor
import os
import threading
from dotenv import load_dotenv
from functools import wraps
from datetime import datetime
import json
import io
import base64
import qrcode
import re
import traceback
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
try:
    from google import genai
except ImportError:
    try:
        import google.generativeai as genai
    except ImportError:
        genai = None
from google_sheets_sync import sync_inventory_to_sheets
from utils import format_history, summarize_history
from threading import Timer
try:
    from docxtpl import DocxTemplate, InlineImage
    from docx.shared import Mm
    from docxcompose.composer import Composer
    from docx import Document
except ImportError as e:
    print(f"[WARNING] Word/docx packages not available: {e}")
    DocxTemplate = InlineImage = Mm = Composer = Document = None

_sync_timer = None
_last_sheets_import = None   # זמן הסנכרון האחרון שיטס→אתר

def trigger_debounced_sync():
    global _sync_timer
    if _sync_timer is not None:
        _sync_timer.cancel()
    # דיליי של 5 שניות כדי לא להציף את גוגל בבקשות
    _sync_timer = Timer(5.0, sync_inventory_to_sheets)
    _sync_timer.start()

# ── AUTO-POLLER: שיטס → אתר כל 3 דקות ─────────────────────────────
def _auto_import_loop():
    """לולאה ברקע: כל 3 דקות מושכת שינויים מגוגל שיטס לאתר"""
    import time
    global _last_sheets_import
    # המתן 30 שניות לפני הריצה הראשונה (לתת ל-Flask להתייצב)
    time.sleep(30)
    while True:
        try:
            from google_sheets_sync import import_from_sheets
            success, msg, stats = import_from_sheets()
            _last_sheets_import = datetime.now()
            if success:
                updated = stats.get('updated', 0)
                if updated > 0:
                    print(f"[AUTO-SYNC] ✅ {_last_sheets_import.strftime('%H:%M:%S')} — עודכנו {updated} רשומות מגיליון שיטס", flush=True)
                else:
                    print(f"[AUTO-SYNC] ⏳ {_last_sheets_import.strftime('%H:%M:%S')} — אין שינויים בשיטס", flush=True)
            else:
                print(f"[AUTO-SYNC] ⚠️ {msg}", flush=True)
        except Exception as ex:
            print(f"[AUTO-SYNC] ❌ שגיאה: {ex}", flush=True)
        time.sleep(180)  # 3 דקות

# הפעל את הפולר בתחילת האפליקציה
_poller_thread = threading.Thread(target=_auto_import_loop, daemon=True, name="SheetsAutoPoller")
_poller_thread.start()
print("[AUTO-SYNC] 🔄 Auto-poller הופעל — יסנכרן שיטס→אתר כל 3 דקות", flush=True)

# טעינת הגדרות
load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv('SECRET_KEY', 'uri_system_2026')

from datetime import timedelta
app.permanent_session_lifetime = timedelta(days=365)
app.config['SESSION_COOKIE_SECURE']   = True
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'

@app.before_request
def refresh_session():
    """מרענן את הסשן בכל בקשה כדי שלא יפוג"""
    if 'user_id' in session:
        session.permanent = True
        session.modified  = True
        # עדכן timestamp פעם ב-5 דקות לכל יותר (למנוע עומס על DB)
        now = datetime.now()
        last_ping = session.get('_last_ping')
        if not last_ping or (now - datetime.fromisoformat(last_ping)).total_seconds() > 300:
            session['_last_ping'] = now.isoformat()
            try:
                conn = get_db_connection()
                if conn:
                    cur = get_safe_cursor(conn)
                    cur.execute("UPDATE users SET timestamp = NOW() WHERE id = %s", (session['user_id'],))
                    conn.commit()
                    cur.close()
                    release_db_connection(conn)
            except Exception:
                pass


# Initialize Google AI (Gemini)
try:
    genai_client = genai.Client(api_key=os.getenv("GOOGLE_API_KEY"))
except Exception as e:
    print(f"[WARNING] Google AI init failed: {e}")
    genai_client = None

# Initialize Connection Pool variables (will be created lazily on first DB access)
db_pool = None
db_pool_initialized = False
IS_LOCAL_MODE = False

# מטמון גלובלי למניעת בדיקה כפולה איטית בגוגל שיטס
attendance_cache = {}


class SafeCursor:
    """Wrapper for cursor to handle %s -> ? translation for SQLite"""
    def __init__(self, cursor, is_sqlite=False):
        self.cursor = cursor
        self.is_sqlite = is_sqlite

    def execute(self, query, params=None):
        if self.is_sqlite and params:
            # Handle list for IN clauses and basic %s replacements
            if "IN (" in query and isinstance(params, (list, tuple)):
                # This is a bit tricky, but common in this app
                pass # Already handled by placeholders in most cases
            query = query.replace('%s', '?')
            # Handle ILIKE -> LIKE for SQLite (SQLite LIKE is case-insensitive usually, but ILIKE is Postgres specific)
            query = query.replace('ILIKE', 'LIKE')
            # Handle NOW() -> datetime('now')
            query = query.replace('NOW()', "datetime('now', 'localtime')")
            # Handle NULLS LAST (SQLite supports it in newer versions, but let's be safe)
            # query = query.replace('NULLS LAST', '') 
        
        try:
            if params:
                return self.cursor.execute(query, params)
            else:
                return self.cursor.execute(query)
        except Exception as e:
            print(f"[DB ERROR] Query: {query}")
            print(f"[DB ERROR] Params: {params}")
            raise e

    def fetchone(self): return self.cursor.fetchone()
    def fetchall(self): return self.cursor.fetchall()
    def close(self): return self.cursor.close()
    def __getattr__(self, name): return getattr(self.cursor, name)

def get_db_connection():
    global db_pool, db_pool_initialized, IS_LOCAL_MODE
    
    if not db_pool_initialized:
        db_url = os.getenv('RENDER_DB_URL') or os.getenv('DATABASE_URL')
        if db_url:
            if 'connect_timeout' not in db_url:
                db_url += ('&' if '?' in db_url else '?') + 'connect_timeout=3'
            try:
                # Fast check: extract host and try resolving to prevent DNS blocking
                import urllib.parse
                parsed = urllib.parse.urlparse(db_url)
                host = parsed.hostname
                if host:
                    import socket
                    # Set a short timeout for socket calls
                    socket.setdefaulttimeout(3)
                    socket.gethostbyname(host)
                
                db_pool = psycopg2.pool.SimpleConnectionPool(1, 15, db_url)
                print("[OK] Database connection pool created successfully (lazy)", flush=True)
                IS_LOCAL_MODE = False
            except Exception as e:
                print(f"[FALLBACK] Cloud DB init failed: {e}. Switching to Local SQLite.", flush=True)
                db_pool = None
                IS_LOCAL_MODE = True
        else:
            db_pool = None
            IS_LOCAL_MODE = True
        db_pool_initialized = True

    if IS_LOCAL_MODE or not db_pool:
        # Fallback to SQLite
        try:
            conn = sqlite3.connect('system_data.db', check_same_thread=False)
            conn.row_factory = sqlite3.Row
            return conn
        except Exception as e:
            print(f"[CRITICAL] SQLite connection failed: {e}", flush=True)
            return None
            
    try:
        conn = db_pool.getconn()
        try:
            # Ping connection to ensure it's alive
            with conn.cursor() as c:
                c.execute('SELECT 1')
        except Exception:
            # Connection is dead, throw it away and get a new one
            db_pool.putconn(conn, close=True)
            conn = db_pool.getconn()
        return conn
    except Exception as e:
        print(f"[FALLBACK] Cloud DB checkout failed: {e}. Switching to Local SQLite.", flush=True)
        try:
            conn = sqlite3.connect('system_data.db', check_same_thread=False)
            conn.row_factory = sqlite3.Row
            return conn
        except Exception as e2:
            print(f"[CRITICAL] All DB connections failed: {e2}")
            return None

def get_safe_cursor(conn):
    if IS_LOCAL_MODE:
        return SafeCursor(conn.cursor(), is_sqlite=True)
    else:
        return conn.cursor(cursor_factory=RealDictCursor)

def release_db_connection(conn):
    if IS_LOCAL_MODE:
        if conn: conn.close()
    elif db_pool and conn:
        db_pool.putconn(conn)


@app.teardown_appcontext
def close_db(error):
    # This ensures connections are released if forgotten, 
    # though it's better to do it manually in routes.
    pass

def run_startup_migrations():
    """הוספת עמודות חדשות למסד אם עדיין לא קיימות"""
    conn = get_db_connection()
    if not conn:
        return
    try:
        cur = get_safe_cursor(conn)
        # הוסף עמודת sheets_delete_request אם לא קיימת
        try:
            if IS_LOCAL_MODE:
                cur.execute("ALTER TABLE computers ADD COLUMN sheets_delete_request INTEGER DEFAULT 0")
            else:
                cur.execute("ALTER TABLE computers ADD COLUMN IF NOT EXISTS sheets_delete_request BOOLEAN DEFAULT FALSE")
            conn.commit()
            print("[OK] Migration: added sheets_delete_request column")
        except Exception:
            conn.rollback()
        # צור טבלת פרויקטים אם לא קיימת
        try:
            if IS_LOCAL_MODE:
                cur.execute("""CREATE TABLE IF NOT EXISTS projects (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    name TEXT NOT NULL,
                    keywords TEXT NOT NULL,
                    sheets_id TEXT DEFAULT '',
                    drive_url TEXT DEFAULT '',
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )""")
            else:
                cur.execute("""CREATE TABLE IF NOT EXISTS projects (
                    id SERIAL PRIMARY KEY,
                    name TEXT NOT NULL,
                    keywords TEXT NOT NULL,
                    sheets_id TEXT DEFAULT '',
                    drive_url TEXT DEFAULT '',
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )""")
            conn.commit()
            print("[OK] Migration: projects table ready")
        except Exception as e:
            conn.rollback()
            print(f"[WARNING] projects table: {e}")
        cur.close()
    finally:
        release_db_connection(conn)

with app.app_context():
    run_startup_migrations()


def get_google_creds(scopes):
    """Load Google credentials from env var (Render) or local file."""
    from google.oauth2.service_account import Credentials
    import json as _json
    sa_json = os.getenv('GOOGLE_SERVICE_ACCOUNT_JSON')
    if sa_json:
        try:
            return Credentials.from_service_account_info(_json.loads(sa_json), scopes=scopes)
        except Exception as e:
            print(f"[CREDS] GOOGLE_SERVICE_ACCOUNT_JSON error: {e}")
    sa_file = os.getenv('GOOGLE_SERVICE_ACCOUNT_FILE', 'service_account.json')
    if os.path.exists(sa_file):
        return Credentials.from_service_account_file(sa_file, scopes=scopes)
    raise FileNotFoundError(f"Google credentials not found. Set GOOGLE_SERVICE_ACCOUNT_JSON on Render.")


@app.context_processor
def utility_processor():
    def get_cage_color(cage):
        if not cage: return "inherit"
        val = sum(ord(c) for c in str(cage))
        hue = (val * 137) % 360
        return f"hsl({hue}, 70%, 65%)"

    def get_computer_spec(computer_number):
        """מחזיר מפרט לפי מספר מחשב"""
        try:
            num = int(str(computer_number).strip())
        except (ValueError, TypeError):
            return None
        if 1 <= num <= 600:
            return {'manufacturer': 'Dell', 'cpu': 'i7-8550U @ 1.80GHz', 'ram': '32GB', 'icon': '🖥️'}
        elif 1001 <= num <= 1600:
            return {'manufacturer': 'HP', 'cpu': 'i5-7200U @ 2.50GHz', 'ram': '8GB', 'icon': '🖥️'}
        elif 2001 <= num <= 2400:
            return {'manufacturer': 'HP', 'cpu': 'i5-7200U @ 2.50GHz', 'ram': '8GB', 'icon': '🖥️'}
        elif 3001 <= num <= 3200:
            return {'manufacturer': 'Dell', 'cpu': 'i7-8550U @ 1.80GHz', 'ram': '16GB', 'icon': '🖥️'}
        elif 4001 <= num <= 4300:
            return {'manufacturer': 'Lenovo', 'cpu': 'i5-7200U @ 2.50GHz', 'ram': '8GB', 'icon': '🖥️'}
        return None

    return dict(get_cage_color=get_cage_color, IS_LOCAL_MODE=IS_LOCAL_MODE, get_computer_spec=get_computer_spec)



@app.template_filter('format_history')
def format_history_filter(val_str):
    return format_history(val_str)

@app.template_filter('summarize_history')
def summarize_history_filter(entry):
    return summarize_history(entry)

@app.template_filter('israel_time')
def israel_time_filter(dt):
    """ממיר datetime מ-UTC לשעון ישראל (UTC+3)"""
    if not dt:
        return '—'
    from datetime import timedelta, timezone
    if hasattr(dt, 'tzinfo') and dt.tzinfo is not None:
        il = dt.astimezone(timezone(timedelta(hours=3)))
    else:
        il = dt + timedelta(hours=3)
    return il.strftime('%H:%M %d/%m/%Y')


def get_auto_spec(barcode):
    """מחזיר מפרט אוטומטי לפי מספר מחשב"""
    try:
        num = int(str(barcode).strip())
    except (ValueError, TypeError):
        return ''
    if 1 <= num <= 600:
        return 'Dell | i7-8550U @ 1.80GHz | 32GB RAM'
    elif 1001 <= num <= 1600:
        return 'HP | i5-7200U @ 2.50GHz | 8GB RAM'
    elif 2001 <= num <= 2400:
        return 'HP | i5-7200U @ 2.50GHz | 8GB RAM'
    elif 3001 <= num <= 3200:
        return 'Dell | i7-8550U @ 1.80GHz | 16GB RAM'
    elif 4001 <= num <= 4300:
        return 'Lenovo | i5-7200U @ 2.50GHz | 8GB RAM'
    return ''


def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user' not in session: return redirect(url_for('login', next=request.url))
        return f(*args, **kwargs)
    return decorated_function

@app.route('/api/projects')
@login_required
def api_projects():
    """מחזיר רשימת פרויקטים + Sheet IDs לסורק"""
    conn = get_db_connection()
    projects = []
    if conn:
        try:
            cur = get_safe_cursor(conn)
            cur.execute("SELECT name, keywords, sheets_id, drive_url FROM projects WHERE sheets_id != '' ORDER BY name")
            rows = cur.fetchall()
            projects = [dict(r) for r in rows]
            cur.close()
        except Exception:
            pass
        finally:
            release_db_connection(conn)
    return jsonify(projects)

@app.route('/')
def index():
    return redirect(url_for('portal')) if 'user' in session else redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '').strip()
        
        print(f"DEBUG: Login attempt - username: '{username}', password: '{password}'")
        
        # בדוק admin_uri מהמסד תחילה (אם קיים שם) — אחרת fallback לקשיח
        admin_matched = False
        if username.lower() in ("uri", "admin_uri") or True:  # always try DB first for any user
            pass  # falls through to DB check below

        # Hardcoded super-admin fallback (only if DB has no custom admin record)
        if username.lower() in ("uri", "admin_uri"):
            conn_check = get_db_connection()
            db_admin = None
            if conn_check:
                try:
                    cur_check = get_safe_cursor(conn_check)
                    cur_check.execute("SELECT username, password FROM users WHERE role='admin' AND username NOT IN ('uri','admin_uri') LIMIT 1")
                    # נסה למצוא admin_uri בDB
                    cur_check.execute("SELECT username, password FROM users WHERE username = %s", (username,))
                    db_admin = cur_check.fetchone()
                    cur_check.close()
                except Exception:
                    pass
                finally:
                    release_db_connection(conn_check)
            
            if not db_admin:
                # fallback hardcoded
                if (username.lower() == "uri" and password == "1234") or (username.lower() == "admin_uri" and password == "uri*"):
                    session.update({
                        'user': username,
                        'user_id': 1,
                        'username': "אורי מנהל מערכת",
                        'role': 'admin'
                    })
                    session.permanent = True
                    print(f"[OK] User {username} logged in (hardcoded fallback)")
                    next_page = request.args.get('next')
                    if next_page:
                        return redirect(next_page)
                    return redirect(url_for('portal'))

        # Check database
        conn = get_db_connection()
        if conn:
            try:
                cur = get_safe_cursor(conn)
                cur.execute("SELECT * FROM users WHERE username = %s", (username,))
                user = cur.fetchone()
                
                if user:
                    db_pass = (user['password'] or '').strip()
                    is_valid = False
                    needs_migration = False
                    
                    # Check if the stored password is a hash (starts with scrypt:, pbkdf2:, bcrypt$, argon2$)
                    if any(db_pass.startswith(prefix) for prefix in ['scrypt:', 'pbkdf2:', 'bcrypt$', 'argon2$']):
                        if check_password_hash(db_pass, password):
                            is_valid = True
                    else:
                        # Plain text match fallback
                        if db_pass == password:
                            is_valid = True
                            needs_migration = True
                    
                    if is_valid:
                        # If plain text, hash and migrate it now
                        if needs_migration:
                            try:
                                hashed_pass = generate_password_hash(password)
                                cur.execute("UPDATE users SET password = %s WHERE id = %s", (hashed_pass, user['id']))
                                conn.commit()
                                print(f"[MIGRATION] Successfully hashed plain-text password for user: {username}", flush=True)
                            except Exception as em:
                                print(f"[MIGRATION ERROR] Could not hash password for user {username}: {em}", flush=True)
                                
                        # Update last active time
                        cur.execute("UPDATE users SET timestamp = NOW() WHERE id = %s", (user['id'],))
                        conn.commit()
                        
                        session.update({
                            'user': user['username'],
                            'user_id': user.get('id', 999), # In case id is missing
                            'username': user['username'],
                            'role': user['role']
                        })
                        session.permanent = True
                        print(f"[OK] User {username} logged in via DB")
                        next_page = request.args.get('next')
                        if next_page:
                            return redirect(next_page)
                        return redirect(url_for('portal'))
                    else:
                        flash("שם משתמש או סיסמה שגויים", "danger")
                else:
                    flash("שם משתמש או סיסמה שגויים", "danger")
            except Exception as e:
                print(f"DB Login Error: {e}")
                flash("שגיאה בהתחברות למסד הנתונים", "danger")
            finally:
                release_db_connection(conn)
        else:
            flash("שגיאת חיבור למסד הנתונים", "danger")
            
    return render_template('login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '').strip()
        
        if not username or not password:
            flash("יש להזין שם משתמש וסיסמה", "warning")
            return redirect(url_for('register'))
            
        conn = get_db_connection()
        if conn:
            try:
                cur = get_safe_cursor(conn)
                cur.execute("SELECT * FROM users WHERE username = %s", (username,))
                if cur.fetchone():
                    flash("שם המשתמש כבר קיים במערכת, בחר שם אחר או התחבר", "warning")
                else:
                    hashed_pass = generate_password_hash(password)
                    cur.execute("INSERT INTO users (username, password, role) VALUES (%s, %s, 'technician')", (username, hashed_pass))
                    conn.commit()
                    flash(f"משתמש {username} נוצר בהצלחה! כעת ניתן להתחבר.", "success")
                    cur.close()
                    return redirect(url_for('login'))
                cur.close()
            except Exception as e:
                print(f"DB Register Error: {e}")
                flash("שגיאה ביצירת המשתמש", "danger")
            finally:
                release_db_connection(conn)
        else:
            flash("שגיאת חיבור למסד הנתונים", "danger")
            
    return render_template('register.html')

@app.route('/portal')
@login_required
def portal():
    return render_template('portal.html')


# ── API: שינוי פרטי כניסה של admin_uri ──────────────────────────────
@app.route('/api/change-admin-credentials', methods=['POST'])
@login_required
def api_change_admin_credentials():
    if session.get('role') != 'admin' and session.get('user') not in ('uri', 'admin_uri'):
        return {"success": False, "error": "גישה מותרת לאדמין בלבד"}, 403

    data = request.json or {}
    new_username = (data.get('new_username') or '').strip()
    new_password = (data.get('new_password') or '').strip()

    if not new_username and not new_password:
        return {"success": False, "error": "לא סופק שום שינוי"}

    conn = get_db_connection()
    if not conn:
        return {"success": False, "error": "DB connection failed"}, 500
    try:
        cur = get_safe_cursor(conn)
        current_username = session.get('user', 'admin_uri')

        # בדוק אם קיים רשומה בDB עבור המשתמש הזה
        cur.execute("SELECT id FROM users WHERE username = %s", (current_username,))
        existing = cur.fetchone()

        if existing:
            # עדכן רשומה קיימת
            if new_username and new_password:
                hashed = generate_password_hash(new_password)
                cur.execute("UPDATE users SET username=%s, password=%s WHERE username=%s",
                            (new_username, hashed, current_username))
            elif new_username:
                cur.execute("UPDATE users SET username=%s WHERE username=%s",
                            (new_username, current_username))
            elif new_password:
                hashed = generate_password_hash(new_password)
                cur.execute("UPDATE users SET password=%s WHERE username=%s",
                            (hashed, current_username))
        else:
            # צור רשומה חדשה במסד עם הפרטים החדשים
            final_username = new_username or current_username
            final_password = new_password or 'uri*'
            hashed = generate_password_hash(final_password)
            cur.execute("INSERT INTO users (username, password, role) VALUES (%s, %s, 'admin')",
                        (final_username, hashed))

        conn.commit()
        cur.close()

        msg = []
        if new_username: msg.append(f"שם משתמש שונה ל-{new_username}")
        if new_password: msg.append("סיסמה עודכנה")
        return {"success": True, "message": " | ".join(msg) + ". מתנתק..."}
    except Exception as e:
        conn.rollback()
        return {"success": False, "error": str(e)}, 500
    finally:
        release_db_connection(conn)


@app.route('/install-cert')
def install_cert():
    """מאפשר לאייפון להוריד ולהתקין את אישור ה-SSL"""
    cert_path = os.path.join(os.path.dirname(__file__), 'server.crt')
    if not os.path.exists(cert_path):
        return "Certificate not found", 404
    return send_file(cert_path, 
                     mimetype='application/x-x509-ca-cert',
                     as_attachment=False,
                     download_name='uri-system.crt')

@app.route('/dashboard')
@app.route('/manage-computers')
@app.route('/computers') # תמיכה בשני השמות
@login_required
def computers():
    search = request.args.get('q', '').strip()
    cage_search = request.args.get('cage_q', '').strip()
    status_filter = request.args.get('status', '').strip()
    page = request.args.get('page', 1, type=int)
    sort = request.args.get('sort', 'scan_time')
    direction = request.args.get('dir', 'desc').lower()
    
    per_page = 100
    offset = (page - 1) * per_page
    
    # Whitelist of allowed sort columns
    allowed_sorts = {
        'barcode': 'barcode',
        'case_number': 'case_number',
        'cage_number': 'cage_number',
        'status': 'status',
        'location': 'location',
        'scan_time': 'scan_time',
        'scan_time': 'scan_time',
        'exam_appeal': 'exam_appeal',
        'specs': 'specs',
        'project': 'project'
    }
    sort_col = allowed_sorts.get(sort, 'scan_time')
    sort_dir = 'ASC' if direction == 'asc' else 'DESC'
    
    conn = get_db_connection()
    if not conn: return "<h1>⚠️ המערכת לא מצליחה להתחבר לענן. בדוק חיבור אינטרנט.</h1>"
    try:
        cur = get_safe_cursor(conn)
        
        # Dashboard Stats
        cur.execute("SELECT COUNT(*) as total FROM computers")
        total_in_db = cur.fetchone()['total']
        
        cur.execute("SELECT status, COUNT(*) as count FROM computers GROUP BY status")
        stats = cur.fetchall()
        stats_dict = {row['status']: row['count'] for row in stats}
        faulty_count = stats_dict.get('תקול', 0)
        # Count computers with no cage assigned
        cur.execute("SELECT COUNT(*) as count FROM computers WHERE (cage_number IS NULL OR TRIM(cage_number) = '') AND (cage_name IS NULL OR TRIM(cage_name) = '')")
        not_in_cage_count = cur.fetchone()['count']
        
        # Base query for computers and count of total matching records
        base_where = " WHERE 1=1"
        params = []
        
        # Free search across multiple fields
        if search:
            # Normalize search term if it's purely digits (barcode-like)
            norm_search = re.sub(r'^0+(?=\d)', '', search)
            base_where += " AND (barcode ILIKE %s OR barcode ILIKE %s OR case_number ILIKE %s OR location ILIKE %s OR notes ILIKE %s OR exam_appeal ILIKE %s)"
            search_val = f"%{search}%"
            norm_val = f"%{norm_search}%"
            params.extend([search_val, norm_val, search_val, search_val, search_val, search_val])
            
        # Dedicated cage search
        if cage_search:
            base_where += " AND (cage_number = %s OR cage_name ILIKE %s)"
            cs = f"%{cage_search}%"
            params.extend([cage_search, cs])
            
        if status_filter:
            base_where += " AND status = %s"
            params.append(status_filter)

        # Get total matching count for pagination
        cur.execute("SELECT COUNT(*) as cnt FROM computers" + base_where, params)
        total_matching = cur.fetchone()['cnt']
        total_pages = (total_matching + per_page - 1) // per_page
            
        # Query results for current page
        query = "SELECT id, barcode, case_number, cage_name, cage_number, location, status, exam_appeal, specs, project, notes, last_technician, scan_time as last_seen FROM computers"
        query += base_where
        
        # Order by logic
        if sort_col == 'scan_time':
            query += f" ORDER BY {sort_col} {sort_dir} NULLS LAST"
        else:
            query += f" ORDER BY {sort_col} {sort_dir}"
            
        query += " LIMIT %s OFFSET %s"
        
        cur.execute(query, params + [per_page, offset])
        computers = cur.fetchall()
        cur.close()
        
        return render_template('computers.html', 
                               computers=computers, 
                               search=search, 
                               status_filter=status_filter,
                               total=total_in_db, 
                               faulty=faulty_count, 
                               not_in_cage=not_in_cage_count,
                               page=page,
                               total_pages=total_pages,
                               total_matching=total_matching,
                               sort=sort,
                               direction=direction,
                               cage_search=cage_search)
    finally:
        release_db_connection(conn)

# נתיבים נוספים שנדרשים בטמפלייט base.html
@app.route('/quick-add', methods=['GET'])
@login_required
def quick_add():
    return render_template('batch_add.html')

@app.route('/add-computer', methods=['GET', 'POST'])
@login_required
def add_computer():
    if request.method == 'POST':
        data = request.form
        conn = get_db_connection()
        if not conn: return "DB connection failed", 500
        try:
            cur = get_safe_cursor(conn)
            barcode = re.sub(r'^0+(?=\d)', '', data['barcode'].strip())
            force_update = data.get('force_update', '0') == '1'

            # Check for existing barcode
            cur.execute("SELECT id, cage_number, location FROM computers WHERE barcode = %s", (barcode,))
            existing = cur.fetchone()

            if existing and not force_update:
                existing_cage = existing['cage_number'] or 'לא ידוע'
                flash(f"⚠️ מחשב {barcode} כבר קיים בכלוב {existing_cage}", "warning")
                return render_template('computer_form.html', action='add', computer=None,
                                       existing_barcode=barcode, existing_cage=existing_cage,
                                       existing_id=existing['id'], prefill=data)

            project = data.get('project', '').strip()
            auto_spec = get_auto_spec(barcode)
            specs_val = data.get('specs', '').strip() or auto_spec

            if existing and force_update:
                # עדכן את הרשומה הקיימת
                cur.execute("""
                    UPDATE computers
                    SET case_number=%s, cage_number=%s, status=%s, location=%s,
                        specs=%s, project=%s, notes=%s, last_technician=%s, scan_time=NOW()
                    WHERE id=%s
                """, (data.get('case_number',''), data.get('cage_number',''),
                       data.get('status','תקין'), data.get('location',''),
                       specs_val, project, data.get('notes',''),
                       session.get('username'), existing['id']))
                cur.execute("""
                    INSERT INTO inventory_history (computer_id, technician, change_type, new_value)
                    VALUES (%s, %s, 'Updated via Add Form', %s)
                """, (existing['id'], session.get('username'),
                       f"כלוב שונה ל: {data.get('cage_number','')}" ))
                conn.commit()
                cur.close()
                flash(f"מחשב {barcode} עודכן בהצלחה!", "success")
                return redirect(url_for('computers'))

            cur.execute("""
                INSERT INTO computers (barcode, case_number, cage_number, status, location, specs, project, notes, scan_time, last_technician)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, NOW(), %s)
            """, (barcode, data.get('case_number',''), data.get('cage_number',''),
                  data.get('status','תקין'), data.get('location',''),
                  specs_val, project, data.get('notes',''), session.get('username')))
            conn.commit()
            cur.close()
            flash("מחשב נוסף בהצלחה!", "success")
            return redirect(url_for('computers'))
        except Exception as e:
            conn.rollback()
            flash(f"שגיאה בהוספת מחשב: {e}", "danger")
        finally:
            release_db_connection(conn)

    return render_template('computer_form.html', action='add', computer=None)

@app.route('/edit-computer/<int:cid>', methods=['GET', 'POST'])
@login_required
def edit_computer(cid):
    conn = get_db_connection()
    if not conn: return "DB connection failed", 500
    try:
        cur = get_safe_cursor(conn)
        if request.method == 'POST':
            data = request.form
            cur.execute("SELECT * FROM computers WHERE id = %s", (cid,))
            old_val = cur.fetchone()
            
            project = data.get('project', '').strip()

            cur.execute("""
                UPDATE computers 
                SET case_number=%s, cage_number=%s, status=%s, location=%s, exam_appeal=%s, specs=%s, project=%s, notes=%s, last_technician=%s
                WHERE id=%s
            """, (data.get('case_number',''), data.get('cage_number',''), data.get('status','תקין'), data.get('location',''), data.get('exam_appeal',''), data.get('specs',''), project, data.get('notes',''), session.get('username'), cid))
            
            cur.execute("""
                INSERT INTO inventory_history (computer_id, technician, change_type, old_value, new_value)
                VALUES (%s, %s, 'Manual Edit', %s, %s)
            """, (
                cid, 
                session.get('username'), 
                json.dumps(dict(old_val), default=str) if old_val else None, 
                json.dumps(dict(data), default=str)
            ))
            
            conn.commit()
            cur.close()
            flash("פרטי המחשב עודכנו!", "success")
            return redirect(url_for('computers'))
            
        cur.execute("SELECT * FROM computers WHERE id = %s", (cid,))
        computer = cur.fetchone()
        cur.close()
        return render_template('computer_form.html', action='edit', computer=computer)
    finally:
        release_db_connection(conn)

@app.route('/delete-computer/<int:cid>', methods=['POST'])
@login_required
def delete_computer(cid):
    conn = get_db_connection()
    if not conn: return "DB connection failed", 500
    try:
        cur = get_safe_cursor(conn)
        # Only admin_uri can hard delete
        if session.get('user') == 'admin_uri':
            cur.execute("DELETE FROM computers WHERE id = %s", (cid,))
            flash("המחשב נמחק מהמערכת סופית (admin_uri)", "warning")
        else:
            cur.execute("UPDATE computers SET status = 'ממתין למחיקה' WHERE id = %s", (cid,))
            flash("הבקשה למחיקת המחשב הועברה לאישור מנהל העל (admin_uri)", "info")
            
        conn.commit()
        cur.close()
    except Exception as e:
        flash(f"שגיאה במחיקה: {e}", "danger")
    finally:
        release_db_connection(conn)
    return redirect(url_for('computers'))

@app.route('/scanner')
@login_required
def scanner():
    return render_template('scanner.html')

@app.route('/exam')
@login_required
def exam_page():
    conn = get_db_connection()
    if not conn: return redirect(url_for('dashboard'))
    try:
        cur = get_safe_cursor(conn)
        cur.execute("SELECT * FROM computers WHERE exam_appeal IS NOT NULL AND TRIM(exam_appeal) != ''")
        computers = cur.fetchall()
        cur.close()
        return render_template('exam.html', computers=computers)
    finally:
        release_db_connection(conn)

import re

@app.route('/history')
@login_required
def history_page():
    conn = get_db_connection()
    if not conn: return redirect(url_for('dashboard'))
    try:
        cur = get_safe_cursor(conn)
        cur.execute("""
            SELECT h.*, c.barcode 
            FROM inventory_history h
            LEFT JOIN computers c ON h.computer_id = c.id
            ORDER BY h.timestamp DESC
            LIMIT 100
        """)
        history = cur.fetchall()
        cur.close()
        
        # Safe processing for templates
        processed_history = []
        for h in history:
            h_dict = dict(h)
            ts = h_dict['timestamp']
            if ts and isinstance(ts, str):
                try:
                    # SQLite default format
                    h_dict['timestamp'] = datetime.strptime(ts, "%Y-%m-%d %H:%M:%S")
                except:
                    pass
            processed_history.append(h_dict)

        return render_template('history.html', history=processed_history)
    finally:
        release_db_connection(conn)

# ── SHARED SCAN SESSION ─────────────────────────────────────────────────────
# מאפשר לטכנאי אחד להגדיר הגדרות, והשאר מצטרפים עם קוד קצר
import random
import string
import time as _time

_shared_sessions = {}  # { code: { settings, owner, created_at } }

def _generate_code():
    """מייצר קוד קצר ייחודי בן 4 תווים"""
    while True:
        code = ''.join(random.choices(string.ascii_uppercase + string.digits, k=4))
        if code not in _shared_sessions:
            return code

def _cleanup_old_sessions():
    """מוחק sessions ישנות (מעל 8 שעות)"""
    now = _time.time()
    expired = [c for c, s in _shared_sessions.items() if now - s['created_at'] > 28800]
    for c in expired:
        del _shared_sessions[c]

@app.route('/api/shared-session/create', methods=['POST'])
@login_required
def create_shared_session():
    """יוצר session חדש עם הגדרות — מחזיר קוד קצר לשיתוף"""
    _cleanup_old_sessions()
    data = request.json or {}
    settings = {
        'location': data.get('location', ''),
        'cage': data.get('cage', ''),
        'status': data.get('status', 'תקין'),
        'exam': data.get('exam', ''),
        'specs': data.get('specs', ''),
        'project': data.get('project', ''),
        'ministry': data.get('ministry', ''),
    }
    code = _generate_code()
    _shared_sessions[code] = {
        'settings': settings,
        'owner': session.get('username', ''),
        'created_at': _time.time()
    }
    print(f"[SharedSession] Created session {code} by {session.get('username')}")
    return {'success': True, 'code': code}

@app.route('/api/shared-session/<code>', methods=['GET'])
@login_required
def get_shared_session(code):
    """מחזיר את הגדרות ה-session לפי קוד"""
    sess = _shared_sessions.get(code.upper())
    if not sess:
        return {'success': False, 'error': 'קוד לא נמצא או פג תוקף'}, 404
    return {'success': True, 'settings': sess['settings'], 'owner': sess['owner']}

@app.route('/api/shared-session/<code>/update', methods=['POST'])
@login_required
def update_shared_session(code):
    """מעדכן הגדרות session קיים (רק הבעלים)"""
    sess = _shared_sessions.get(code.upper())
    if not sess:
        return {'success': False, 'error': 'קוד לא נמצא'}, 404
    if sess['owner'] != session.get('username'):
        return {'success': False, 'error': 'רק הבעלים יכול לעדכן'}, 403
    data = request.json or {}
    for key in ['location', 'cage', 'status', 'exam', 'specs', 'project', 'ministry']:
        if key in data:
            sess['settings'][key] = data[key]
    return {'success': True}

@app.route('/api/shared-session/<code>/close', methods=['POST'])
@login_required
def close_shared_session(code):
    """סוגר session"""
    code = code.upper()
    if code in _shared_sessions:
        if _shared_sessions[code]['owner'] == session.get('username'):
            del _shared_sessions[code]
    return {'success': True}
# ────────────────────────────────────────────────────────────────────────────

@app.route('/api/process-scan', methods=['POST'])
@login_required
def process_scan():
    barcode = request.json.get('barcode', '').strip()
    if not barcode:
        return {"error": "No barcode provided"}, 400
    barcode = re.sub(r'^0+(?=\d)', '', barcode)

    # ── חסימת QR קודים של נבחנים ──────────────────────────────────────
    # מזהים: מכיל | (pipe) עם נתוני בחינה, מילות מפתח של מבחן/ערעור/נבחן
    EXAM_KEYWORDS = ['במבחן', 'ערעור', 'נבחן', 'בחינה', 'תעודה']
    is_exam_qr = (
        '|' in barcode and len(barcode) > 20  # מבנה QR נבחן: data|data|data
        or any(kw in barcode for kw in EXAM_KEYWORDS)
    )
    if is_exam_qr:
        print(f"[BLOCKED] Examinee QR rejected: {barcode[:30]}...")
        return {
            "error": "❌ QR זה שייך לנבחן — לא ניתן לסרוק אותו כמחשב!",
            "blocked": True
        }, 400
    # ──────────────────────────────────────────────────────────────────

    # ── ולידציה: ברקוד חייב להיות מספרי בלבד ובטווח 1-9999 ──────────
    if not barcode.isdigit() or not (1 <= int(barcode) <= 9999):
        print(f"[BLOCKED] Invalid barcode rejected: '{barcode}'")
        return {
            "error": "❌ הברקוד דפוק",
            "blocked": True
        }, 400
    # ──────────────────────────────────────────────────────────────────

    conn = get_db_connection()
    if not conn: return {"error": "DB connection failed"}, 500
    try:
        cur = get_safe_cursor(conn)
        # Check if computer exists
        cur.execute("SELECT * FROM computers WHERE barcode = %s ORDER BY id DESC LIMIT 1", (barcode,))
        computer = cur.fetchone()
        
        if computer:
            # Update the existing record instead of inserting a duplicate
            auto_spec = get_auto_spec(barcode)
            cur.execute("""
                UPDATE computers 
                SET scan_time = NOW(), last_technician = %s
                    {}
                WHERE id = %s
            """.format(", specs = %s" if auto_spec else ""),
                (session.get('username'),) + ((auto_spec, computer['id']) if auto_spec else (computer['id'],)))
            
            # Fetch the updated record (SQLite doesn't support RETURNING)
            cur.execute("SELECT * FROM computers WHERE id = %s", (computer['id'],))
            new_computer = dict(cur.fetchone())
            
            # Fetch previous technician and time (skip current scan = LIMIT 1 OFFSET 1)
            cur.execute("SELECT technician, timestamp FROM inventory_history WHERE computer_id = %s ORDER BY timestamp DESC LIMIT 1 OFFSET 1", (new_computer['id'],))
            hist = cur.fetchone()
            new_computer['last_technician'] = hist['technician'] if hist and hist['technician'] else ""
            
            ts = hist['timestamp'] if hist else None
            if ts and not isinstance(ts, str):
                new_computer['last_scan_time'] = ts.strftime("%d/%m/%Y %H:%M")
            else:
                new_computer['last_scan_time'] = str(ts) if ts else ''
            
            # Serialize scan_time for JSON
            if new_computer.get('scan_time') and not isinstance(new_computer['scan_time'], str):
                new_computer['scan_time'] = new_computer['scan_time'].strftime("%d/%m/%Y %H:%M")
            
            conn.commit()
            cur.close()
            # Trigger Google Sheets sync in background
            trigger_debounced_sync()
            
            # מחזירים את המידע הקיים כדי שהטופס יתמלא נכון
            return {"exists": True, "computer": new_computer}
        else:
            # Create completely new record
            auto_spec = get_auto_spec(barcode)
            cur.execute("""
                INSERT INTO computers (barcode, status, scan_time, specs, notes, last_technician) 
                VALUES (%s, 'תקין', NOW(), %s, %s, %s) 
            """, (barcode, auto_spec or None, None, session.get('username')))
            
            last_id = cur.lastrowid if hasattr(cur, 'lastrowid') else None
            if last_id:
                cur.execute("SELECT * FROM computers WHERE id = %s", (last_id,))
                new_computer = dict(cur.fetchone())
            else:
                # Fallback if lastrowid fails
                cur.execute("SELECT * FROM computers WHERE barcode = %s ORDER BY id DESC LIMIT 1", (barcode,))
                new_computer = dict(cur.fetchone())

            conn.commit()
            cur.close()
            # Trigger Google Sheets sync in background
            trigger_debounced_sync()
            
            return {"exists": False, "computer": new_computer}
            
    except Exception as e:
        print(f"Error in process_scan: {e}")
        return {"error": str(e)}, 500
    finally:
        release_db_connection(conn)

@app.route('/api/update-computer', methods=['POST'])
@login_required
def api_update_computer():
    data = request.json
    cid = data.get('id')
    if not cid: return {"error": "No ID provided"}, 400

    conn = get_db_connection()
    if not conn: return {"error": "DB connection failed"}, 500
    try:
        cur = get_safe_cursor(conn)
        # Get old values for history
        cur.execute("SELECT * FROM computers WHERE id = %s", (cid,))
        old_val = cur.fetchone()
        
        # Update
        updates = []
        params = []
        for key in ['case_number', 'cage_number', 'status', 'location', 'exam_appeal', 'specs', 'project', 'notes']:
            if key in data:
                val = data[key]
                updates.append(f"{key} = %s")
                params.append(val)
        
        # Always update last_technician on scan update
        updates.append("last_technician = %s")
        params.append(session.get('username'))

        if updates:
            params.append(cid)
            cur.execute(f"UPDATE computers SET {', '.join(updates)} WHERE id = %s", params)
            
            # Record history
            cur.execute("""
                INSERT INTO inventory_history (computer_id, technician, change_type, old_value, new_value)
                VALUES (%s, %s, 'Update via Scan', %s, %s)
            """, (
                cid, 
                session.get('username'), 
                json.dumps(dict(old_val), default=str) if old_val else None, 
                json.dumps(data, default=str)
            ))
            
            conn.commit()
            cur.close()
            # Trigger Google Sheets sync in background
            trigger_debounced_sync()
            
            return {"success": True}
        return {"success": False, "message": "No fields to update"}
        
    except Exception as e:
        print(f"Error in api_update_computer: {e}")
        return {"error": str(e)}, 500
    finally:
        release_db_connection(conn)

# ── AI ASSISTANT (GEMINI) ───────────────────────────────────────────
@app.route('/api/ai-chat', methods=['POST'])
@login_required
def api_ai_chat():
    data = request.json
    user_msg = data.get('message', '').strip()
    if not user_msg:
        return {"error": "No message provided"}, 400

    conn = get_db_connection()
    if not conn: return {"error": "DB connection failed"}, 500
    try:
        cur = get_safe_cursor(conn)
        # Fetch stats for context
        cur.execute("SELECT COUNT(*) as total FROM computers")
        total = cur.fetchone()['total']
        
        cur.execute("SELECT status, COUNT(*) as count FROM computers GROUP BY status")
        status_stats = cur.fetchall()
        status_desc = ", ".join([f"{row['status']}: {row['count']}" for row in status_stats])
        
        # Build system context
        system_prompt = f"""
        אתה עוזר ה-AI של מערכת URI לניהול מלאי מחשבים. 
        הנתונים הנוכחיים במערכת הם:
        - סה"כ מחשבים: {total}
        - סטטוסים: {status_desc}
        
        ענה למשתמש בעברית בצורה עוזרת, מקצועית וקצרה. 
        אם המשתמש שואל על המצב, השתמש בנתונים שלעיל.
        משתמש נוכחי: {session.get('username')}
        """
        
        response = genai_client.models.generate_content(
            model='gemini-2.0-flash',
            contents=[system_prompt, user_msg]
        )
        
        return {"response": response.text}
        
    except Exception as e:
        print(f"AI Error: {e}")
        return {"error": str(e)}, 500
    finally:
        release_db_connection(conn)

# ── API: Batch Operations ──────────────────────────────────────────────
@app.route('/api/batch-update', methods=['POST'])
@login_required
def api_batch_update():
    data = request.json
    ids = data.get('ids', [])
    updates = data.get('updates', {})
    
    if not ids or not updates:
        return {"success": False, "error": "Missing ids or updates"}, 400
        
    conn = get_db_connection()
    if not conn: return {"success": False, "error": "DB connection failed"}, 500
    try:
        cur = get_safe_cursor(conn)
        
        set_clauses = []
        params = []
        for key in ['location', 'cage_number', 'cage_name', 'status', 'exam_appeal', 'specs', 'project', 'ministry', 'notes']:
            if key in updates:
                set_clauses.append(f"{key} = %s")
                params.append(updates[key])
                
        if not set_clauses:
            return {"success": False, "error": "No valid fields provided"}, 400
            
        params.extend(ids)
        placeholders = ','.join(['%s'] * len(ids))
        
        query = f"UPDATE computers SET {', '.join(set_clauses)} WHERE id IN ({placeholders})"
        cur.execute(query, params)
        
        # Log to history for each (simplified to avoid mass select first, assuming identical change)
        for cid in ids:
            cur.execute("""
                INSERT INTO inventory_history (computer_id, technician, change_type, new_value)
                VALUES (%s, %s, 'Batch Update', %s)
            """, (cid, session.get('username'), json.dumps(updates, default=str)))
            
        conn.commit()
        cur.close()
        return {"success": True}
    except Exception as e:
        print(f"Error in batch-update: {e}")
        return {"success": False, "error": str(e)}, 500
    finally:
        release_db_connection(conn)

@app.route('/api/batch-delete', methods=['POST'])
@login_required
def api_batch_delete():
    # Only admin should be able to trigger this in the UI, verify on server too
    if session.get('role') != 'admin' and session.get('user') != 'admin_uri':
        return {"success": False, "error": "Unauthorized"}, 403
        
    data = request.json
    ids = data.get('ids', [])
    if not ids: return {"success": False, "error": "No ids provided"}, 400
    
    conn = get_db_connection()
    if not conn: return {"success": False, "error": "DB connection failed"}, 500
    try:
        cur = get_safe_cursor(conn)
        placeholders = ','.join(['%s'] * len(ids))
        cur.execute(f"DELETE FROM computers WHERE id IN ({placeholders})", ids)
        conn.commit()
        cur.close()
        return {"success": True}
    except Exception as e:
        print(f"Error in batch-delete: {e}")
        return {"success": False, "error": str(e)}, 500
    finally:
        release_db_connection(conn)


# ── API: Google Sheets Sync ──────────────────────────────────────────
@app.route('/api/sync-to-sheets', methods=['POST'])
@login_required
def api_sync_to_sheets():
    """
    Manual trigger for Google Sheets sync.
    """
    # Only admin or experienced technicians should sync
    # For now, allow all logged in users as requested "connect table"
    success, message = sync_inventory_to_sheets()
    if success:
        return {"success": True, "message": message}
    else:
        return {"success": False, "error": message}, 500


@app.route('/api/find-duplicates', methods=['GET'])
@login_required
def api_find_duplicates():
    conn = get_db_connection()
    if not conn:
        return jsonify({"success": False, "error": "DB error"}), 500
    try:
        cur = get_safe_cursor(conn)
        cur.execute("""
            SELECT computer_number, COUNT(*) as cnt
            FROM computers
            GROUP BY computer_number
            HAVING COUNT(*) > 1
            ORDER BY cnt DESC
        """)
        rows = cur.fetchall()
        duplicates = []
        for r in rows:
            if hasattr(r, 'keys'):
                duplicates.append({'computer_number': r['computer_number'], 'count': r['cnt']})
            else:
                duplicates.append({'computer_number': r[0], 'count': r[1]})
        cur.close()
        return jsonify({"success": True, "duplicates": duplicates, "total": len(duplicates)})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        release_db_connection(conn)


@app.route('/api/delete-duplicates', methods=['POST'])
@login_required
def api_delete_duplicates():
    if session.get('role') != 'admin' and session.get('user') != 'admin_uri':
        return jsonify({"success": False, "error": "Admin only"}), 403
    conn = get_db_connection()
    if not conn:
        return jsonify({"success": False, "error": "DB error"}), 500
    try:
        cur = get_safe_cursor(conn)
        cur.execute("""
            DELETE FROM computers
            WHERE id NOT IN (
                SELECT DISTINCT ON (computer_number) id
                FROM computers
                ORDER BY computer_number, last_scan DESC NULLS LAST
            )
        """)
        deleted = cur.rowcount
        conn.commit()
        cur.close()
        return jsonify({"success": True, "deleted": deleted, "message": f"נמחקו {deleted} כפולים"})
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        release_db_connection(conn)


# ── API: ייבוא מגוגל שיטס → מסד (רק admin) ──────────────────────────
@app.route('/api/import-from-sheets', methods=['POST'])
@login_required
def api_import_from_sheets():
    if session.get('role') != 'admin' and session.get('user') != 'admin_uri':
        return {"success": False, "error": "גישה מותרת לאדמין בלבד"}, 403
    from google_sheets_sync import import_from_sheets
    success, message, stats = import_from_sheets()
    return {"success": success, "message": message, "stats": stats}


# ── API: אישור מחיקה סופית (רק admin) ───────────────────────────────
@app.route('/api/approve-sheets-delete', methods=['POST'])
@login_required
def api_approve_sheets_delete():
    """רק admin_uri יכול לאשר מחיקה סופית של מחשבים שמסומנים כ-sheets_delete_request"""
    if session.get('user') != 'admin_uri' and session.get('role') != 'admin':
        return {"success": False, "error": "גישה מותרת לאדמין בלבד"}, 403

    data = request.json or {}
    action = data.get('action')  # 'approve' or 'cancel'
    computer_ids = data.get('ids', [])

    if not computer_ids:
        # אם לא נשלחו IDs — פעל על כולם המסומנים
        conn = get_db_connection()
        if not conn:
            return {"success": False, "error": "DB connection failed"}, 500
        try:
            cur = get_safe_cursor(conn)
            cur.execute("SELECT id FROM computers WHERE sheets_delete_request = TRUE")
            rows = cur.fetchall()
            computer_ids = [r['id'] for r in rows]
            cur.close()
        finally:
            release_db_connection(conn)

    if not computer_ids:
        return {"success": True, "message": "אין מחשבים הממתינים לאישור מחיקה", "count": 0}

    conn = get_db_connection()
    if not conn:
        return {"success": False, "error": "DB connection failed"}, 500
    try:
        cur = get_safe_cursor(conn)
        placeholders = ','.join(['%s'] * len(computer_ids))

        if action == 'approve':
            # מחיקה סופית
            cur.execute(f"DELETE FROM computers WHERE id IN ({placeholders}) AND sheets_delete_request = TRUE", computer_ids)
            conn.commit()
            count = cur.rowcount
            cur.close()
            return {"success": True, "message": f"נמחקו {count} מחשבים לצמיתות", "count": count}
        else:
            # ביטול סימון — המחשב נשאר במסד
            cur.execute(f"UPDATE computers SET sheets_delete_request = FALSE WHERE id IN ({placeholders})", computer_ids)
            conn.commit()
            count = cur.rowcount
            cur.close()
            return {"success": True, "message": f"בוטל סימון המחיקה עבור {count} מחשבים", "count": count}
    except Exception as e:
        return {"success": False, "error": str(e)}, 500
    finally:
        release_db_connection(conn)


# ── API: שליפת מידע כלוב ──────────────────────────────────────────────
@app.route('/api/cage/<cage_id>', methods=['GET'])
@login_required
def api_get_cage(cage_id):
    """מחזיר מידע על כלוב + רשימת המחשבים בו"""
    conn = get_db_connection()
    if not conn: return {"error": "DB connection failed"}, 500
    try:
        cur = get_safe_cursor(conn)

        # שליפת פרטי הכלוב (אם קיים בטבלת cages)
        cur.execute("SELECT * FROM cages WHERE cage_id = %s", (cage_id,))
        cage = cur.fetchone()

        # שליפת מחשבים בכלוב זה
        cur.execute("""
            SELECT id, barcode, status, location, specs, scan_time, notes
            FROM computers
            WHERE cage_number = %s OR cage_name = %s
            ORDER BY scan_time DESC NULLS LAST
        """, (cage_id, cage_id))
        computers_in_cage = cur.fetchall()

        # סטטיסטיקות
        total = len(computers_in_cage)
        status_counts = {}
        for c in computers_in_cage:
            s_val = c.get('status')
            s = s_val if s_val is not None else 'לא ידוע'
            status_counts[s] = status_counts.get(s, 0) + 1  # type: ignore

        cur.close()
        return {
            "cage": dict(cage) if cage else {"cage_id": cage_id, "name": f"כלוב {cage_id}"},
            "computers": [dict(c) for c in computers_in_cage],
            "total": total,
            "status_counts": status_counts
        }
    except Exception as e:
        print(f"Error in api_get_cage: {e}")
        return {"error": str(e)}, 500
    finally:
        release_db_connection(conn)

# ── API: סריקה מהירה (ללא חלון) ───────────────────────────────────────
@app.route('/api/fast-scan', methods=['POST'])
@login_required
def api_fast_scan():
    """
    סריקה מהירה - מעדכן מחשב ישירות ללא תצוגת UI מפורטת.
    user שולח: barcode, location, cage_number, cage_name, status
    """
    data = request.json
    barcode = data.get('barcode', '').strip()
    if not barcode:
        return {"success": False, "error": "ברקוד חסר"}, 400
    barcode = re.sub(r'^0+(?=\d)', '', barcode)

    location   = data.get('location', '')
    cage_number = data.get('cage_number', '')
    cage_name  = data.get('cage_name', cage_number)
    status     = data.get('status', 'תקין')
    specs      = data.get('specs', '') or get_auto_spec(barcode)
    project    = data.get('project', '')
    ministry   = data.get('ministry', '')
    technician = session.get('username', 'לא ידוע')

    conn = get_db_connection()
    if not conn: return {"success": False, "error": "DB connection failed"}, 500
    try:
        cur = get_safe_cursor(conn)
        cur.execute("SELECT * FROM computers WHERE barcode = %s ORDER BY id DESC LIMIT 1", (barcode,))
        computer = cur.fetchone()

        old_val = dict(computer) if computer else None
        last_technician = "לא ידוע"
        notes_val = ""

        if old_val:
            notes_val = old_val.get('notes') or ""
            fields_to_update = []
            params = []
            
            for key in ['location', 'cage_number', 'cage_name', 'status', 'specs', 'project', 'ministry', 'exam_appeal']:
                if key in data:
                    fields_to_update.append(f"{key} = %s")
                    params.append(data[key])
            
            # Special auto-ministry logic
            if 'project' in data and 'ministry' not in data:
                project_name = data['project']
                ministry = ''
                if 'רופאי' in project_name or 'שיניים' in project_name: ministry = 'משרד הבריאות'
                elif 'משפטים' in project_name: ministry = 'משרד המשפטים'
                elif 'עבודה' in project_name: ministry = 'משרד העבודה'
                elif 'חינוך' in project_name: ministry = 'משרד החינוך'
                fields_to_update.append("ministry = %s")
                params.append(ministry)
                
            fields_to_update.extend(["scan_time = NOW()", "last_technician = %s"])
            params.extend([technician, old_val['id']])
            
            query = f"UPDATE computers SET {', '.join(fields_to_update)} WHERE id = %s RETURNING *"
            cur.execute(query, params)
            new_computer = cur.fetchone()
            
            cur.execute("SELECT technician, timestamp FROM inventory_history WHERE computer_id = %s ORDER BY timestamp DESC LIMIT 1", (new_computer['id'],))
            hist = cur.fetchone()
            if hist and hist['technician']:
                last_technician = hist['technician']
                last_scan_time = hist['timestamp'].strftime("%d/%m/%Y %H:%M") if hist['timestamp'] else ""
        else:
            fields = ['barcode', 'scan_time', 'last_technician']
            params = [barcode, technician]
            
            for key in ['location', 'cage_number', 'cage_name', 'status', 'specs', 'project', 'ministry', 'exam_appeal']:
                if key in data:
                    fields.append(key)
                    params.append(data[key])
                    
            if 'project' in data and 'ministry' not in data:
                fields.append('ministry')
                project_name = data['project']
                ministry = ''
                if 'רופאי' in project_name or 'שיניים' in project_name: ministry = 'משרד הבריאות'
                elif 'משפטים' in project_name: ministry = 'משרד המשפטים'
                elif 'עבודה' in project_name: ministry = 'משרד העבודה'
                elif 'חינוך' in project_name: ministry = 'משרד החינוך'
                params.append(ministry)
                
            if 'status' not in data:
                fields.append('status')
                params.append('תקין')
                
            placeholders = ', '.join(['%s'] * len(fields))
            query = f"INSERT INTO computers ({', '.join(fields)}) VALUES ({placeholders}) RETURNING *"
            cur.execute(query, params)
            new_computer = cur.fetchone()

        cur.execute("""
            INSERT INTO inventory_history (computer_id, technician, change_type, old_value, new_value)
            VALUES (%s, %s, 'Fast Scan', %s, %s)
        """, (
            new_computer['id'],
            technician,
            json.dumps(old_val, default=str) if old_val else None,
            json.dumps(data, default=str)
        ))

        conn.commit()
        cur.close()
        
        # טריגר לסנכרון אוטומטי (מתוזמן) עבור סריקה מהירה
        trigger_debounced_sync()
        
        return {
            "success": True, 
            "barcode": barcode, 
            "is_new": old_val is None,
            "previous_cage": old_val.get('cage_number', '') if old_val else "",
            "notes": notes_val
        }

    except Exception as e:
        print(f"Error in api_fast_scan: {e}")
        return {"success": False, "error": str(e)}, 500
    finally:
        release_db_connection(conn)


@app.route('/api/admin_approve_delete', methods=['POST'])
@login_required
def admin_approve_delete():
    # Only admin_uri can approve deletes
    if session.get('user') != 'admin_uri':
        return {"success": False, "error": "Unauthorized: Only admin_uri can perform this action"}, 403
        
    data = request.json
    barcode = data.get('barcode')
    action = data.get('action')
    
    if not barcode or not action:
        return {"success": False, "error": "Missing parameters"}, 400
        
    conn = get_db_connection()
    if not conn: return {"success": False, "error": "DB Error"}, 500
    try:
        cur = get_safe_cursor(conn)
        if action == 'hard_delete':
            cur.execute("DELETE FROM computers WHERE barcode = %s", (barcode,))
        elif action == 'restore':
            cur.execute("UPDATE computers SET status = 'תקין' WHERE barcode = %s", (barcode,))
            
        conn.commit()
        cur.close()
        return {"success": True}
    except Exception as e:
        print(f"Error in admin_approve_delete: {e}")
        return {"success": False, "error": str(e)}, 500
    finally:
        release_db_connection(conn)

@app.route('/manage-users')
@login_required
def manage_users():
    if session.get('role') != 'admin' and session.get('user') != 'admin_uri':
        flash("אין לך הרשאה לגשת לעמוד זה", "danger")
        return redirect(url_for('portal'))
        
    conn = get_db_connection()
    if not conn: return "DB connection failed", 500
    try:
        cur = get_safe_cursor(conn)
        # 1. Fetch users
        cur.execute("SELECT username, role, timestamp FROM users ORDER BY timestamp DESC NULLS LAST")
        users_raw = cur.fetchall()

        # קבע מי מחובר (פעיל ב-30 דקות האחרונות)
        now = datetime.now()
        users = []
        for u in users_raw:
            ts = u['timestamp'] if hasattr(u, '__getitem__') else u[2]
            is_online = False
            if ts:
                if hasattr(ts, 'replace'):
                    ts_naive = ts.replace(tzinfo=None) if hasattr(ts, 'tzinfo') and ts.tzinfo else ts
                    is_online = (now - ts_naive).total_seconds() < 1800  # 30 דקות
            users.append({'username': u['username'] if hasattr(u, '__getitem__') else u[0],
                          'role': u['role'] if hasattr(u, '__getitem__') else u[1],
                          'timestamp': ts,
                          'is_online': is_online})

        # 2. Fetch pending deletions
        cur.execute("SELECT barcode as computer_number, cage_number FROM computers WHERE status = 'ממתין למחיקה'")
        pending_raw = cur.fetchall()

        # Format pending for template
        pending = []
        for p in pending_raw:
            pending.append({
                'computer_number': p['computer_number'],
                'cage_number': p['cage_number'] or 'לא ידוע',
                'scanned_by': 'טכנאי'
            })

        cur.close()
        return render_template('manage_users.html', users=users, pending=pending)
    finally:
        release_db_connection(conn)

@app.route('/api/add_user', methods=['POST'])
@login_required
def api_add_user():
    # Only admin should be able to add users
    if session.get('role') != 'admin' and session.get('user') != 'admin_uri':
        return {"success": False, "error": "Unauthorized"}, 403
        
    data = request.json
    username = data.get('username', '').strip()
    password = data.get('password', '').strip()
    role = data.get('role', 'technician')
    
    if not username or not password:
        return {"success": False, "error": "Missing username or password"}, 400
        
    conn = get_db_connection()
    if not conn: return {"success": False, "error": "DB connection failed"}, 500
    try:
        cur = get_safe_cursor(conn)
        hashed_pass = generate_password_hash(password)
        cur.execute("INSERT INTO users (username, password, role) VALUES (%s, %s, %s)", (username, hashed_pass, role))
        conn.commit()
        cur.close()
        return {"success": True}
    except Exception as e:
        print(f"Error adding user: {e}")
        return {"success": False, "error": str(e)}, 500
    finally:
        release_db_connection(conn)

@app.route('/api/update_user', methods=['POST'])
@login_required
def api_update_user():
    if session.get('role') != 'admin' and session.get('user') not in ('uri', 'admin_uri'):
        return {"success": False, "error": "Unauthorized"}, 403
    data = request.json or {}
    orig_username = data.get('orig_username', '').strip()
    new_username  = data.get('new_username', '').strip()
    new_password  = (data.get('password') or '').strip()
    new_role      = data.get('role', '').strip()
    allowed_roles = ['technician', 'scanner', 'manager', 'logistics', 'admin']
    if not orig_username or not new_username:
        return {"success": False, "error": "שם משתמש לא תקין"}, 400
    if new_role and new_role not in allowed_roles:
        return {"success": False, "error": "הרשאה לא תקינה"}, 400
    if orig_username == 'admin_uri':
        return {"success": False, "error": "לא ניתן לערוך admin_uri מכאן"}, 400
    conn = get_db_connection()
    if not conn: return {"success": False, "error": "DB connection failed"}, 500
    try:
        cur = get_safe_cursor(conn)
        if new_password:
            hashed = generate_password_hash(new_password)
            cur.execute("UPDATE users SET username=%s, password=%s, role=%s WHERE username=%s",
                        (new_username, hashed, new_role or 'technician', orig_username))
        else:
            cur.execute("UPDATE users SET username=%s, role=%s WHERE username=%s",
                        (new_username, new_role or 'technician', orig_username))
        conn.commit()
        cur.close()
        return {"success": True}
    except Exception as e:
        conn.rollback()
        return {"success": False, "error": str(e)}, 500
    finally:
        release_db_connection(conn)


@app.route('/api/update_user_password', methods=['POST'])
@login_required
def api_update_user_password():
    if session.get('role') != 'admin' and session.get('user') not in ('uri', 'admin_uri'):
        return {"success": False, "error": "Unauthorized"}, 403
    data = request.json or {}
    username = data.get('username', '').strip()
    new_password = data.get('password', '').strip()
    if not username or len(new_password) < 4:
        return {"success": False, "error": "נתונים לא תקינים"}, 400
    conn = get_db_connection()
    if not conn: return {"success": False, "error": "DB connection failed"}, 500
    try:
        cur = get_safe_cursor(conn)
        hashed = generate_password_hash(new_password)
        cur.execute("UPDATE users SET password = %s WHERE username = %s", (hashed, username))
        conn.commit()
        cur.close()
        return {"success": True}
    except Exception as e:
        return {"success": False, "error": str(e)}, 500
    finally:
        release_db_connection(conn)


@app.route('/api/update_user_role', methods=['POST'])
@login_required
def api_update_user_role():
    if session.get('role') != 'admin' and session.get('user') not in ('uri', 'admin_uri'):
        return {"success": False, "error": "Unauthorized"}, 403
    data = request.json or {}
    username = data.get('username', '').strip()
    new_role = data.get('role', '').strip()
    allowed_roles = ['technician', 'scanner', 'manager', 'logistics', 'admin']
    if not username or new_role not in allowed_roles:
        return {"success": False, "error": "נתונים לא תקינים"}, 400
    if username == 'admin_uri':
        return {"success": False, "error": "לא ניתן לשנות את הרשאות admin_uri"}, 400
    conn = get_db_connection()
    if not conn: return {"success": False, "error": "DB connection failed"}, 500
    try:
        cur = get_safe_cursor(conn)
        cur.execute("UPDATE users SET role = %s WHERE username = %s", (new_role, username))
        conn.commit()
        cur.close()
        return {"success": True}
    except Exception as e:
        return {"success": False, "error": str(e)}, 500
    finally:
        release_db_connection(conn)


@app.route('/api/delete_user/<username>', methods=['DELETE'])
@login_required
def api_delete_user(username):
    # Only admin should be able to delete users
    if session.get('role') != 'admin' and session.get('user') != 'admin_uri':
        return {"success": False, "error": "Unauthorized"}, 403
        
    if username == 'admin_uri':
        return {"success": False, "error": "Cannot delete super-admin"}, 400
        
    conn = get_db_connection()
    if not conn: return {"success": False, "error": "DB connection failed"}, 500
    try:
        cur = get_safe_cursor(conn)
        cur.execute("DELETE FROM users WHERE username = %s", (username,))
        conn.commit()
        cur.close()
        return {"success": True}
    except Exception as e:
        print(f"Error deleting user: {e}")
        return {"success": False, "error": str(e)}, 500
    finally:
        release_db_connection(conn)

@app.route('/cages')
@login_required
def cages_page():
    conn = get_db_connection()
    if not conn: return "DB Error", 500
    try:
        cur = get_safe_cursor(conn)
        cur.execute("""
            SELECT
                c.cage_id, c.name, c.location, c.notes,
                COUNT(comp.id) AS computer_count,
                SUM(CASE WHEN comp.status = 'תקין' THEN 1 ELSE 0 END) AS ok_count,
                SUM(CASE WHEN comp.status = 'תקול' THEN 1 ELSE 0 END) AS broken_count,
                SUM(CASE WHEN comp.status NOT IN ('תקין','תקול') AND comp.status IS NOT NULL THEN 1 ELSE 0 END) AS other_count
            FROM cages c
            LEFT JOIN computers comp ON comp.cage_number = c.cage_id OR comp.cage_name = c.cage_id
            GROUP BY c.id, c.cage_id, c.name, c.location, c.notes
            ORDER BY computer_count DESC
        """)
        cages = cur.fetchall()
        cur.close()
        return render_template('cages.html', cages=cages)
    except Exception as e:
        return f"<h1>Error: {e}</h1>", 500
    finally:
        release_db_connection(conn)

@app.route('/api/cage/save', methods=['POST'])
@login_required
def api_save_cage():
    data = request.json
    cage_id = data.get('cage_id', '').strip()
    existing_id = data.get('existing_id', '').strip()
    if not cage_id: return {'success': False, 'error': 'cage_id is required'}, 400
    conn = get_db_connection()
    if not conn: return {'success': False, 'error': 'DB Error'}, 500
    try:
        cur = get_safe_cursor(conn)
        if existing_id:
            cur.execute("UPDATE cages SET name=%s, location=%s, notes=%s, updated_at=NOW() WHERE cage_id=%s", 
                        (data.get('name',''), data.get('location',''), data.get('notes',''), existing_id))
        else:
            cur.execute("INSERT INTO cages (cage_id, name, location, notes) VALUES (%s, %s, %s, %s) ON CONFLICT (cage_id) DO UPDATE SET name=EXCLUDED.name, location=EXCLUDED.location, notes=EXCLUDED.notes, updated_at=NOW()", 
                        (cage_id, data.get('name',''), data.get('location',''), data.get('notes','')))
        conn.commit()
        cur.close()
        return {'success': True}
    finally: release_db_connection(conn)

@app.route('/api/pack-cage-photo', methods=['POST'])
@login_required
def api_pack_cage_photo():
    data = request.json
    cage_id = data.get('cage_id', '').strip()
    image_b64 = data.get('image', '')
    
    if not cage_id or not image_b64:
        return {"success": False, "error": "Missing cage_id or image"}, 400
        
    # Extract base64 part
    if ',' in image_b64:
        image_b64 = image_b64.split(',', 1)[1]
        
    try:
        image_data = base64.b64decode(image_b64)
    except Exception as e:
        return {"success": False, "error": "Invalid image data"}, 400

    conn = get_db_connection()
    if not conn: return {"success": False, "error": "DB connection failed"}, 500
    cur = None
    try:
        # Call Gemini Vision
        prompt = "Look at the handwritten numbers written in white on the edges of the laptops in the cage. Extract all of them. Return ONLY a JSON array of strings (e.g. [\"1064\", \"366\", \"1480\"]). Do not add any markdown, comments, or other text."
        response = genai_client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[
                genai.types.Part.from_bytes(data=image_data, mime_type="image/jpeg"),
                prompt
            ]
        )
        
        try:
            # Clean response text just in case Gemini adds markdown
            text = response.text.strip()
            if text.startswith('```json'): text = text[7:]
            if text.startswith('```'): text = text[3:]
            if text.endswith('```'): text = text[:-3]
            text = text.strip()
            extracted_numbers = json.loads(text)
            if not isinstance(extracted_numbers, list):
                extracted_numbers = []
        except Exception as e:
            print(f"Gemini parse error: {e}. Raw response: {response.text}")
            return {"success": False, "error": "Could not parse AI response as JSON list."}, 500
            
        cur = get_safe_cursor(conn)
        
        results = {
            "success_count": 0,
            "transferred_count": 0,
            "new_count": 0,
            "details": []
        }
        
        technician = session.get('username', 'לא ידוע')
        
        # Process each extracted number
        for barcode in extracted_numbers:
            barcode = str(barcode).strip()
            barcode = re.sub(r'^0+(?=\d)', '', barcode)
            if not barcode: continue
            
            cur.execute("SELECT * FROM computers WHERE barcode = %s ORDER BY id DESC LIMIT 1", (barcode,))
            computer = cur.fetchone()
            
            if computer:
                old_val = dict(computer)
                prev_cage = old_val.get('cage_number', '')
                
                # Check if it's already in this cage
                is_transfer = (prev_cage and prev_cage != cage_id)
                
                # Update
                cur.execute("""
                    UPDATE computers 
                    SET cage_number = %s, cage_name = %s, scan_time = NOW(), last_technician = %s
                    WHERE id = %s
                """, (cage_id, cage_id, technician, old_val['id']))
                
                # Log history
                cur.execute("""
                    INSERT INTO inventory_history (computer_id, technician, change_type, old_value, new_value)
                    VALUES (%s, %s, 'Photo Pack', %s, %s)
                """, (
                    old_val['id'],
                    technician,
                    json.dumps(old_val, default=str),
                    json.dumps({"cage_number": cage_id, "cage_name": cage_id}, default=str)
                ))
                
                if is_transfer:
                    results["transferred_count"] += 1
                    results["details"].append({"barcode": barcode, "status": "transferred", "prev_cage": prev_cage})
                else:
                    results["success_count"] += 1
                    results["details"].append({"barcode": barcode, "status": "updated"})
                    
            else:
                # New computer
                cur.execute("""
                    INSERT INTO computers (barcode, cage_number, cage_name, status, scan_time, last_technician)
                    VALUES (%s, %s, %s, 'תקין', NOW(), %s) RETURNING id
                """, (barcode, cage_id, cage_id, technician))
                new_id = cur.fetchone()['id']
                
                cur.execute("""
                    INSERT INTO inventory_history (computer_id, technician, change_type, old_value, new_value)
                    VALUES (%s, %s, 'Photo Pack (New)', %s, %s)
                """, (
                    new_id,
                    technician,
                    None,
                    json.dumps({"cage_number": cage_id, "cage_name": cage_id, "status": "תקין"}, default=str)
                ))
                
                results["new_count"] += 1
                results["details"].append({"barcode": barcode, "status": "new"})
                
        conn.commit()
        cur.close()
        
        # Trigger async sync
        trigger_debounced_sync()
        
        results["total_extracted"] = len(extracted_numbers)
        results["success"] = True
        return results

    except Exception as e:
        if conn:
            conn.rollback()
        print(f"Error in api_pack_cage_photo: {e}")
        traceback.print_exc()
        return {"success": False, "error": str(e)}, 500
    finally:
        if cur:
            try:
                cur.close()
            except Exception:
                pass
        release_db_connection(conn)

@app.route('/scan-dashboard')
@login_required
def scan_dashboard():
    return render_template('scan_dashboard.html')

@app.route('/api/scan-stats')
@login_required
def api_scan_stats():
    conn = get_db_connection()
    if not conn: return {'error': 'DB Error'}, 500
    try:
        cur = get_safe_cursor(conn)
        cur.execute("SELECT COUNT(*) as cnt FROM inventory_history WHERE timestamp::date = CURRENT_DATE AND change_type IN ('Fast Scan', 'Update via Scan')")
        today_total = cur.fetchone()['cnt']
        cur.execute("SELECT COUNT(*) as cnt FROM computers")
        total_computers = cur.fetchone()['cnt']
        cur.execute("SELECT COUNT(*) as cnt FROM computers WHERE status = 'תקול'")
        broken = cur.fetchone()['cnt']
        cur.execute("SELECT technician, COUNT(*) as count, MAX(timestamp) as last_scan FROM inventory_history WHERE timestamp::date = CURRENT_DATE AND change_type IN ('Fast Scan', 'Update via Scan') GROUP BY technician ORDER BY count DESC")
        workers = [dict(r) for r in cur.fetchall()]
        cur.close()
        return {'today_total': today_total, 'total_computers': total_computers, 'broken': broken, 'workers': workers}
    finally: release_db_connection(conn)

@app.route('/export/computers')
@login_required
def export_computers():
    conn = get_db_connection()
    if not conn: return 'DB Error', 500
    try:
        cur = get_safe_cursor(conn)
        cur.execute("SELECT barcode, cage_number, cage_name, location, status, case_number, exam_appeal, notes, scan_time FROM computers ORDER BY cage_number, barcode")
        rows = cur.fetchall()
        cur.close()
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = 'מחשבים'
        ws.sheet_view.rightToLeft = True
        headers = ['ברקוד', 'כלוב', 'שם כלוב', 'מיקום', 'סטטוס', 'מספר תיק', 'מבחן/ערעור', 'הערות', 'נסרק לאחרונה']
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_num, value=header)
            cell.font = Font(bold=True)
        for row_num, row in enumerate(rows, 2):
            values = [row.get('barcode',''), row.get('cage_number',''), row.get('cage_name',''), row.get('location',''), row.get('status',''), row.get('case_number',''), row.get('exam_appeal',''), row.get('notes',''), str(row.get('scan_time',''))[0:16] if row.get('scan_time') else '']
            for col_num, val in enumerate(values, 1): ws.cell(row=row_num, column=col_num, value=val)
        buf = io.BytesIO(); wb.save(buf); buf.seek(0)
        return send_file(buf, as_attachment=True, download_name=f"inventory_{datetime.now().strftime('%Y%m%d')}.xlsx", mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    finally: release_db_connection(conn)

# ══════════════════════════════════════════════════════════════
# מערכת נוכחות נבחנים
# ══════════════════════════════════════════════════════════════

@app.route('/exam-attendance')
@login_required
def exam_attendance():
    """דשבורד נוכחות נבחנים"""
    conn = get_db_connection()
    if not conn: return "DB Error", 500
    try:
        cur = get_safe_cursor(conn)
        cur.execute("SELECT * FROM examinees ORDER BY exam_name, full_name")
        examinees = [dict(e) for e in cur.fetchall()]
        cur.execute("SELECT COUNT(*) as total FROM examinees")
        total = cur.fetchone()['total']
        cur.execute("SELECT COUNT(*) as attended FROM examinees WHERE is_present = 1")
        attended = cur.fetchone()['attended']
        cur.execute("SELECT DISTINCT exam_name FROM examinees WHERE exam_name IS NOT NULL AND exam_name != ''")
        exams = [r['exam_name'] for r in cur.fetchall()]
        cur.close()
        return render_template('exam_attendance.html',
                               examinees=examinees,
                               total=total,
                               attended=attended,
                               not_attended=total - attended,
                               exams=exams)
    finally:
        release_db_connection(conn)

@app.route('/exam-attendance/add', methods=['GET', 'POST'])
@login_required
def exam_attendance_add():
    """הוספת נבחן ידנית"""
    if request.method == 'POST':
        data = request.form
        conn = get_db_connection()
        if not conn: return "DB Error", 500
        try:
            cur = get_safe_cursor(conn)
            import uuid
            token = uuid.uuid4().hex
            cur.execute("""
                INSERT INTO examinees (full_name, id_number, username, password, classroom, exam_name, laptop_number, token)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                data.get('name','').strip(),
                data.get('id_number','').strip(),
                data.get('username','').strip(),
                data.get('password','').strip(),
                data.get('location','').strip(),
                data.get('exam_name','').strip(),
                data.get('computer','').strip(),
                token
            ))
            conn.commit()
            cur.close()
            flash(f"נבחן {data.get('name','')} נוסף בהצלחה! ✅", "success")
            return redirect(url_for('exam_attendance'))
        except Exception as e:
            flash(f"שגיאה: {e}", "danger")
        finally:
            release_db_connection(conn)
    return render_template('exam_attendance_add.html')

@app.route('/exam-attendance/import', methods=['POST'])
@login_required
def exam_attendance_import():
    """ייבוא נבחנים מ-Excel"""
    file = request.files.get('excel_file')
    if not file or not file.filename.endswith(('.xlsx', '.xls')):
        flash("יש להעלות קובץ Excel (.xlsx / .xls)", "danger")
        return redirect(url_for('exam_attendance'))

    try:
        wb = openpyxl.load_workbook(file)
        ws = wb.active
        
        # קריאת כותרת הבחינה משורה 1
        exam_title_row1 = ''
        for cell in ws[1]:
            if cell.value and str(cell.value).strip():
                exam_title_row1 = str(cell.value).strip()
                break

        headers = []
        header_row_idx = 1
        for row_idx, r in enumerate(ws.iter_rows(min_row=1, max_row=10, values_only=True), start=1):
            row_strs = [str(c).strip() if c else '' for c in r]
            if any('שם' in s or 'תעודת' in s or 'ת.ז' in s or 'id' in s.lower() or 'פרטי' in s for s in row_strs):
                headers = row_strs
                header_row_idx = row_idx
                break
        if not headers:
            headers = [str(cell.value).strip() if cell.value else '' for cell in ws[1]]

        # מיפוי עמודות גמיש — כולל שם פרטי + שם משפחה
        col_map = {}
        for i, h in enumerate(headers):
            h_lower = h.lower()
            if 'פרטי' in h:                                          col_map['first_name'] = i
            elif 'משפחה' in h or 'family' in h_lower:               col_map['last_name']  = i
            elif ('שם' in h or 'name' in h_lower) and 'משתמש' not in h: col_map['name'] = i
            elif 'זהות' in h or 'ת.ז' in h or 'id' in h_lower:     col_map['id_number']  = i
            elif 'משתמש' in h or 'user' in h_lower or 'קוד' in h:  col_map['username']   = i
            elif 'סיסמ' in h or 'pass' in h_lower:                  col_map['password']   = i
            elif 'טור' in h or 'עמודה' in h:                        col_map['row_number'] = i
            elif 'כסא' in h or 'כיסא' in h or 'seat' in h_lower or 'מושב' in h:   col_map['seat_number']= i
            elif 'מיקום' in h or 'כיתה' in h or 'location' in h_lower or 'אולם' in h: col_map['location']= i
            elif 'בחינה' in h or 'exam' in h_lower:                 col_map['exam_name']  = i
            elif 'מחשב' in h or 'computer' in h_lower:              col_map['computer']   = i
            elif 'התאמות' in h or 'notes' in h_lower:               col_map['notes']      = i

        def safe_val(val):
            """המרת ערך תא — מטפל במספרים גדולים/נוטציה מדעית"""
            if val is None: return ''
            if isinstance(val, float):
                if val == int(val): return str(int(val))
                return str(val)
            return str(val).strip()

        # איסוף כל השורות
        all_rows = []
        for row in ws.iter_rows(min_row=header_row_idx + 1, values_only=True):
            if not any(row): continue
            def get_val(key, r=row):
                idx = col_map.get(key)
                if idx is None: return ''
                return safe_val(r[idx])

            # שם מלא: אם יש שם פרטי + משפחה — מאחד
            first = get_val('first_name')
            last  = get_val('last_name')
            if first or last:
                name = (first + ' ' + last).strip()
            else:
                name = get_val('name')
            if not name: continue

            all_rows.append({
                'name':        name,
                'id_number':   get_val('id_number'),
                'username':    get_val('username'),
                'password':    get_val('password'),
                'location':    get_val('location'),
                'exam_name':   get_val('exam_name'),
                'computer':    get_val('computer'),
                'notes':       get_val('notes'),
                'row_number':  get_val('row_number'),
                'seat_number': get_val('seat_number'),
            })

        # מיון לפי טור → כסא
        def sort_key(r):
            try:    row_n  = int(float(r['row_number']))  if r['row_number']  else 9999
            except: row_n  = 9999
            try:    seat_n = int(float(r['seat_number'])) if r['seat_number'] else 9999
            except: seat_n = 9999
            return (row_n, seat_n)
        all_rows.sort(key=sort_key)

        # ── Google Drive: תיקייה + גיליון ייעודי לכל מבחן ──
        from drive_manager import create_folder_and_sheet_if_not_exists
        import re

        raw_name = os.path.splitext(file.filename)[0]

        # שם המשרד/בחינה — ניקוי תאריך מהשם
        if exam_title_row1:
            title_clean = re.sub(r'[-–]\s*נוכחות\s*$', '', exam_title_row1).strip()
            title_clean = re.sub(r'[-–]\s*\d+\.\d+\.\d+\s*$', '', title_clean).strip()
            exam_sheet_name = title_clean or raw_name
        else:
            exam_sheet_name = re.sub(r'\s*\d+[\./]\d+[\./]\d+\s*$', '', raw_name).strip() or raw_name

        print(f"[IMPORT] Creating/finding Drive sheet for: '{exam_sheet_name}'", flush=True)

        # יצירת תיקייה + גיליון אוטומטית ב-Drive (אם לא קיים — יוצר, אם קיים — מחזיר)
        new_sheet_id = create_folder_and_sheet_if_not_exists(exam_sheet_name)

        if not new_sheet_id:
            # Fallback לגיליון ברירת המחדל
            new_sheet_id = os.getenv('EXAM_ATTENDANCE_SHEET_ID', '1YWLJA5T8Uq7IGzlzXSA1PPwrdSIPx9eazEcwWXwh3uM')
            print(f"[IMPORT] Fallback to default sheet", flush=True)

        import gspread
        from google.oauth2.service_account import Credentials
        scopes  = ['https://www.googleapis.com/auth/spreadsheets', 'https://www.googleapis.com/auth/drive']
        sa_file = os.getenv('GOOGLE_SERVICE_ACCOUNT_FILE', 'service_account.json')
        creds   = Credentials.from_service_account_file(sa_file, scopes=scopes)
        gs_client = gspread.authorize(creds)
        sh = gs_client.open_by_key(new_sheet_id)
        ws_tab = sh.sheet1

        # ── כתוב נתונים לגיליון — לפי מבנה האקסל ──
        header_row = [
            'שם פרטי', 'שם משפחה', 'ת.ז', 'התאמות', 'סיסמה',
            'שם משתמש', 'גרסה', "סה' אולם", 'טור', 'כסא',
            'מ.מחשב', 'נוכחות', 'שעת סריקה', 'טכנאי'
        ]
        rows_to_write = [header_row]
        for rec in all_rows:
            rows_to_write.append([
                rec['name'],         # שם פרטי (שם מלא)
                '',                  # שם משפחה
                rec['id_number'],    # ת.ז
                rec.get('notes',''), # התאמות
                rec['password'],     # סיסמה
                rec['username'],     # שם משתמש
                rec.get('exam_name', exam_sheet_name),  # גרסה
                '',                  # סה' אולם
                rec['row_number'],   # טור
                rec['seat_number'],  # כסא
                '',                  # מ.מחשב ← ימולא בסריקה
                '',                  # נוכחות ← ימולא בסריקה
                '',                  # שעת סריקה ← ימולא בסריקה
                '',                  # טכנאי ← ימולא בסריקה
            ])
        ws_tab.clear()
        ws_tab.update('A1', rows_to_write, value_input_option='USER_ENTERED')

        # עיצוב כותרות
        last_col = chr(ord('A') + len(header_row) - 1)
        ws_tab.format(f'A1:{last_col}1', {
            'textFormat': {'bold': True},
            'backgroundColor': {'red': 0.8, 'green': 0.9, 'blue': 1.0},
            'horizontalAlignment': 'CENTER'
        })

        # ── שמור מיפוי פרויקט במסד ──
        sheet_url = f"https://docs.google.com/spreadsheets/d/{new_sheet_id}"
        conn2 = get_db_connection()
        if conn2:
            try:
                cur2 = get_safe_cursor(conn2)
                cur2.execute("SELECT id FROM projects WHERE name = %s", (exam_sheet_name,))
                existing = cur2.fetchone()
                if existing:
                    cur2.execute("UPDATE projects SET sheets_id=%s, drive_url=%s WHERE name=%s",
                                 (new_sheet_id, sheet_url, exam_sheet_name))
                else:
                    cur2.execute("INSERT INTO projects (name, keywords, sheets_id, drive_url) VALUES (%s,%s,%s,%s)",
                                 (exam_sheet_name, exam_sheet_name, new_sheet_id, sheet_url))
                conn2.commit()
                cur2.close()
            except Exception:
                conn2.rollback()
            finally:
                release_db_connection(conn2)

        flash(f"✅ {len(all_rows)} נבחנים יובאו! 📂 גיליון נוצר: {exam_sheet_name}", "success")
        flash(f"🔗 <a href='{sheet_url}' target='_blank'>פתח את הגיליון ב-Google Drive</a>", "info")

    except Exception as e:
        import traceback; traceback.print_exc()
        flash(f"שגיאה בייבוא: {e}", "danger")
    return redirect(url_for('exam_attendance'))




@app.route('/api/generate-word-docs', methods=['POST'])
@login_required
def generate_word_docs():
    """מחולל דפי נבחנים מקובץ אקסל ותבנית וורד"""
    excel_file = request.files.get('excel_file')
    word_template = request.files.get('word_template')
    
    if not excel_file:
        flash("יש להעלות קובץ Excel", "danger")
        return redirect(url_for('exam_attendance'))
        
    try:
        # Load Excel data
        wb = openpyxl.load_workbook(excel_file)
        ws = wb.active
        
        # קריאת שם הבחינה משורה 1 (כותרת הדף) אם קיים
        exam_title_from_header = ''
        first_row = next(ws.iter_rows(min_row=1, max_row=1, values_only=True))
        for cell in first_row:
            if cell and str(cell).strip():
                exam_title_from_header = str(cell).strip()
                break

        headers = []
        header_row_idx = 1
        # זיהוי שורת כותרות: צריך לפחות 2 עמודות מוכרות (מונע שגיאות כמו "חשמלאים"⊃"שם")
        HEADER_KEYS = ['שם', 'ת.ז', 'תעודת', 'סיסמ', 'משתמש', 'טור', 'כסא', 'מחשב', 'התאמות', 'פרטי', 'משפחה', 'id', 'pass', 'user', 'seat', 'name']
        for row_idx, r in enumerate(ws.iter_rows(min_row=1, max_row=10, values_only=True), start=1):
            row_strs = [str(c).strip() if c else '' for c in r]
            matches = sum(1 for s in row_strs if any(k in s.lower() for k in HEADER_KEYS))
            if matches >= 2:
                headers = row_strs
                header_row_idx = row_idx
                break
        if not headers:
            headers = [str(cell.value).strip() if cell.value else '' for cell in ws[1]]

        def safe_val(val):
            """המרת ערך תא — מטפל במספרים גדולים/נוטציה מדעית"""
            if val is None: return ''
            if isinstance(val, float):
                if val == int(val): return str(int(val))
                return str(val)
            return str(val).strip()

        col_map = {}
        for i, h in enumerate(headers):
            h_lower = h.lower()
            if 'פרטי' in h:                                              col_map['first_name'] = i
            elif 'משפחה' in h or 'family' in h_lower:                   col_map['last_name']  = i
            elif ('שם' in h or 'name' in h_lower) and 'משתמש' not in h: col_map['name']       = i
            elif 'זהות' in h or 'ת.ז' in h or 'id' in h_lower:         col_map['id_number']  = i
            elif 'משתמש' in h or 'user' in h_lower or 'קוד' in h:      col_map['username']   = i
            elif 'סיסמ' in h or 'pass' in h_lower:                      col_map['password']   = i
            elif 'מיקום' in h or 'כיתה' in h or 'location' in h_lower or 'אולם' in h: col_map['location']   = i
            elif 'בחינה' in h or 'exam' in h_lower:                     col_map['exam_name']  = i
            elif 'מחשב' in h or 'computer' in h_lower:                  col_map['computer']   = i
            elif 'התאמות' in h or 'notes' in h_lower:                   col_map['notes']      = i
            elif 'טור' in h or 'עמודה' in h:                            col_map['row']        = i
            elif 'כסא' in h or 'כיסא' in h or 'seat' in h_lower or 'מושב' in h:       col_map['seat']       = i

        examinees = []
        for row in ws.iter_rows(min_row=header_row_idx + 1, values_only=True):
            if not any(row): continue
            def get_val(key, r=row):
                idx = col_map.get(key)
                if idx is None: return ''
                return safe_val(r[idx])

            # DEBUG — שומר מיפוי לקובץ (למניעת שגיאת charmap בטרמינל)
            if len(examinees) == 0:
                try:
                    with open('debug_colmap.txt', 'w', encoding='utf-8') as dbf:
                        dbf.write(f"col_map = {col_map}\n")
                        dbf.write(f"headers = {headers}\n")
                        dbf.write(f"first data row = {[str(v) for v in row]}\n")
                except Exception:
                    pass

            # שם מלא: שם פרטי + שם משפחה אם קיימים
            first = get_val('first_name')
            last  = get_val('last_name')
            if first or last:
                name = (first + ' ' + last).strip()
            else:
                name = get_val('name')
            if not name: continue
            
            examinees.append({
                'name': name,
                'id_number': get_val('id_number'),
                'username': get_val('username'),
                'password': get_val('password'),
                'location': get_val('location'),
                'row': get_val('row'),
                'seat': get_val('seat'),
                'exam_name': get_val('exam_name'),
                'computer': get_val('computer'),
                'notes': get_val('notes')
            })
            
        if not examinees:
            flash("לא נמצאו נתונים בקובץ האקסל", "warning")
            return redirect(url_for('exam_attendance'))
            
        # Load template
        if word_template and word_template.filename.endswith('.docx'):
            word_bytes = word_template.read()
        else:
            template_path = os.path.join(os.path.dirname(__file__), 'default_template.docx')
            with open(template_path, 'rb') as f:
                word_bytes = f.read()
        
        all_docs = []  # רשימה של (שם_קובץ, bytes) לכל נבחן
        
        for i, e in enumerate(examinees):
            # Load template first so we can use it for InlineImage
            tpl = DocxTemplate(io.BytesIO(word_bytes))
            
            # Format: exam_name|id_number|name|username|password|row|seat
            # exam_name: from column or from Excel title row
            exam_title = e['exam_name'] if e['exam_name'] else (exam_title_from_header or 'EXAMINEE')
            # Clean trailing "- נוכחות" or "נוכחות"
            exam_title_clean = re.sub(r'[-–]\s*נוכחות\s*$', '', exam_title).strip()
            exam_title_clean = exam_title_clean.replace('|', ' ').strip()  # מונע שבירת פורמט ה-QR
            qr_data = f"{exam_title_clean}|{e['id_number']}|{e['name']}|{e['username']}|{e['password']}|{e['row']}|{e['seat']}"
            # QR מ-API מבוטל — משתמשים רק ב-QR מקומי (ראה Step 2 בהמשך)
            qr_inline = ""
            
            # Context
            context = {
                'full_name': e['name'],
                'id_number': e['id_number'],
                'username': e['username'],
                'password': e['password'],
                'location': e['location'],
                'row': e['row'],
                'seat': e['seat'],
                'exam_name': e['exam_name'],
                'computer': e['computer'],
                'notes': e['notes'],
                'qr_code': qr_inline
            }
            
            tpl.render(context)
            
            # Save rendered to memory
            rendered_buf = io.BytesIO()
            tpl.save(rendered_buf)
            rendered_buf.seek(0)
            
            doc = Document(rendered_buf)

            # ---- Step 0.5: Replace static old title and location to match the imported exam ----
            def replace_static_text(p, new_title, new_loc):
                txt = p.text
                title_keywords = ["משרד הבריאות", "מינהל האחיות", "מומחיות", "טיפול תומך", "רישוי חשמלאים", "חשמלאים", "כיתות חשמל", "כיתת חשמל", "חשמל"]
                loc_keywords = ["בניין REIT1", "פתח תקווה", "מליאה", "אפעל 6", "בניין", "קומה", "כיתה"]
                
                is_title = any(kw in txt for kw in title_keywords) and "טופס התחברות" not in txt
                is_loc = any(kw in txt for kw in loc_keywords) and not is_title and "טופס התחברות" not in txt
                
                if is_title:
                    if p.runs:
                        p.runs[0].text = new_title
                        for r_item in p.runs[1:]:
                            r_item.text = ""
                    else:
                        p.text = new_title
                elif is_loc and new_loc:
                    if p.runs:
                        p.runs[0].text = new_loc
                        for r_item in p.runs[1:]:
                            r_item.text = ""
                    else:
                        p.text = new_loc

            # Process sections (headers and footers)
            for section in doc.sections:
                # 1. Process headers with both heuristic and keywords
                for hf_name in ['header', 'first_page_header', 'even_page_header']:
                    hf = getattr(section, hf_name, None)
                    if hf:
                        # Heuristic index-based replacement for direct paragraphs in header
                        non_empty_ps = [p for p in hf.paragraphs if p.text.strip()]
                        title_p = None
                        loc_p = None
                        for p in non_empty_ps:
                            if "טופס התחברות" not in p.text:
                                if not title_p:
                                    title_p = p
                                elif not loc_p:
                                    loc_p = p
                                    break
                        if title_p:
                            if title_p.runs:
                                title_p.runs[0].text = exam_title_clean
                                for r in title_p.runs[1:]:
                                    r.text = ""
                            else:
                                title_p.text = exam_title_clean
                        if loc_p and e['location']:
                            if loc_p.runs:
                                loc_p.runs[0].text = e['location']
                                for r in loc_p.runs[1:]:
                                    r.text = ""
                            else:
                                loc_p.text = e['location']
                        
                        # Also apply keyword replacement for paragraphs and tables inside header
                        for p in hf.paragraphs:
                            replace_static_text(p, exam_title_clean, e['location'])
                        for table in hf.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for p in cell.paragraphs:
                                        replace_static_text(p, exam_title_clean, e['location'])

                # 2. Process footers with keywords replacement only (no heuristic to prevent corruption)
                for hf_name in ['footer', 'first_page_footer', 'even_page_footer']:
                    hf = getattr(section, hf_name, None)
                    if hf:
                        for p in hf.paragraphs:
                            replace_static_text(p, exam_title_clean, e['location'])
                        for table in hf.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for p in cell.paragraphs:
                                        replace_static_text(p, exam_title_clean, e['location'])

            # Process main body paragraphs and tables
            for p in doc.paragraphs:
                replace_static_text(p, exam_title_clean, e['location'])
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            replace_static_text(p, exam_title_clean, e['location'])
            
            # Generate QR locally
            import qrcode as qrcode_lib
            qr_gen = qrcode_lib.QRCode(version=1, box_size=10, border=1)
            qr_gen.add_data(qr_data)
            qr_gen.make(fit=True)
            qr_img = qr_gen.make_image(fill_color="black", back_color="white")
            qr_buf_local = io.BytesIO()
            qr_img.save(qr_buf_local, format='PNG')
            qr_buf_local.seek(0)

            from docx.shared import Inches, Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            # ---- Step 1: Fill data tables by label matching ----
            label_to_value = {
                'שם נבחן': e['name'],
                'תעודת זהות': e['id_number'],
                'קוד משתמש': e['username'],
                'סיסמה': e['password'],
                'התאמות': e['notes'],
            }
            for t_idx, t in enumerate(doc.tables):
                # דלג על טבלת "למילוי על ידי הנבחן/ת" — נשארת ריקה לכתב יד
                table_text = ' '.join(cell.text for row in t.rows for cell in row.cells)
                if 'חתימה' in table_text or 'למילוי' in table_text:
                    continue
                for row_idx, r in enumerate(t.rows):
                    if len(r.cells) >= 2:
                        label_right = r.cells[1].text.strip()
                        label_left  = r.cells[0].text.strip()
                        for key, val_text in label_to_value.items():
                            if key in label_right:
                                r.cells[0].paragraphs[0].clear()
                                r.cells[0].paragraphs[0].add_run(val_text)
                                break
                            elif key in label_left:
                                r.cells[1].paragraphs[0].clear()
                                r.cells[1].paragraphs[0].add_run(val_text)
                                break

            # ---- Step 2: Find seat paragraph in BODY and add QR ----
            seat_display = e.get('seat', '') or str(i + 1)  # fallback: מספר סידורי
            target_p = None
            for p in doc.paragraphs:
                txt = p.text.strip()
                if txt == seat_display or txt == '1':   # '1' הוא ה-placeholder בתבנית
                    target_p = p
                    break

            if target_p:
                # Force LTR on this specific paragraph (so QR is left, number is right)
                pPr = target_p._p.get_or_add_pPr()
                for bidi_el in pPr.findall(qn('w:bidi')):
                    pPr.remove(bidi_el)
                bidi = OxmlElement('w:bidi')
                bidi.set(qn('w:val'), '0')
                pPr.append(bidi)

                target_p.text = ''
                target_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # 1. Add QR Code first (will be on the LEFT in LTR)
                qr_buf_local.seek(0)
                run_qr = target_p.add_run()
                run_qr.add_picture(qr_buf_local, width=Inches(1.2))

                # 2. Add a RIGHT tab stop to push the number to the far right
                tab_stops = target_p.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Cm(16), WD_TAB_ALIGNMENT.RIGHT)
                target_p.add_run('\t')

                # 3. Add Seat Number (will be on the RIGHT)
                run_seat = target_p.add_run(seat_display)
                run_seat.font.size = Pt(85)
                run_seat.font.bold = True
                run_seat.font.name = 'Tahoma'


            # DEBUG: שמור כל מסמך נפרד לבדיקה
            import os as _os
            _debug_dir = r'C:\Users\uri\OneDrive\Desktop\test\debug_docs'
            _os.makedirs(_debug_dir, exist_ok=True)
            _safe = e['name'].replace(' ','_')[:20]
            _dbuf = io.BytesIO()
            doc.save(_dbuf)
            _dbuf.seek(0)
            with open(f'{_debug_dir}\\{i+1:02d}_{_safe}.docx','wb') as _f:
                _f.write(_dbuf.read())
            print(f"[DEBUG] Saved person {i+1}: {e['name']} | qr_data={qr_data[:60]}")

            # שמור כל דוק ברשימה
            all_docs.append(doc)

        # מזג עם docxcompose — מטפל נכון בקשרי תמונות
        from docxcompose.composer import Composer
        from docx.enum.text import WD_BREAK

        master_doc = all_docs[0]
        composer = Composer(master_doc)

        for doc in all_docs[1:]:
            # הוסף מעבר עמוד בסוף המסמך הנוכחי לפני הצירוף
            last_para = master_doc.paragraphs[-1] if master_doc.paragraphs else master_doc.add_paragraph()
            last_para.add_run().add_break(WD_BREAK.PAGE)
            composer.append(doc)

        final_buf = io.BytesIO()
        master_doc.save(final_buf)
        final_buf.seek(0)

        return send_file(
            final_buf,
            as_attachment=True,
            download_name='טפסי_נבחנים.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

        
    except Exception as e:
        import traceback
        traceback.print_exc()
        flash(f"שגיאה ביצירת המסמכים: {e}", "danger")
        return redirect(url_for('exam_attendance'))

@app.route('/exam-attendance/print')
@login_required
def exam_attendance_print():
    """דף הדפסת טפסים עם QR"""
    exam_filter = request.args.get('exam', '')
    conn = get_db_connection()
    if not conn: return "DB Error", 500
    try:
        cur = get_safe_cursor(conn)
        if exam_filter:
            cur.execute("SELECT * FROM examinees WHERE exam_name = %s ORDER BY full_name", (exam_filter,))
        else:
            cur.execute("SELECT * FROM examinees ORDER BY exam_name, full_name")
        examinees = [dict(e) for e in cur.fetchall()]
        cur.close()
        # יצירת QR לכל נבחן
        for e in examinees:
            qr_data = f"EXAMINEE|{e['id_number']}|{e['full_name']}|{e.get('username','')}|{e.get('password','')}|{e.get('classroom','')}|{e.get('exam_name','')}|{e.get('laptop_number','')}"  
            qr_img = qrcode.make(qr_data)
            buf = io.BytesIO()
            qr_img.save(buf)
            buf.seek(0)
            e['qr_b64'] = base64.b64encode(buf.read()).decode('utf-8')
        return render_template('exam_print.html', examinees=examinees, exam_filter=exam_filter)
    finally:
        release_db_connection(conn)

@app.route('/test-scanner', methods=['GET'])
@login_required
def test_scanner():
    return render_template('test_scanner.html')

@app.route('/simple-scanner', methods=['GET'])
@login_required
def simple_scanner():
    """עמוד סורק פשוט - סורק רק נבחנים"""
    return render_template('simple_scanner.html')

@app.route('/api/simple-scan', methods=['POST'])
@login_required
def api_simple_scan():
    data = request.json
    qr_text = data.get('qr', '').strip()
    if not qr_text.startswith('EXAMINEE|'):
        return {"error": "QR לא תקין"}, 400

    parts = qr_text.split('|')
    id_number = parts[1] if len(parts) > 1 else ''
    name = parts[2] if len(parts) > 2 else ''

    conn = get_db_connection()
    if not conn: return {"error": "DB Error"}, 500
    try:
        cur = get_safe_cursor(conn)
        cur.execute("SELECT * FROM examinees WHERE id_number = %s", (id_number,))
        if not cur.fetchone():
            return {"error": f"נבחן עם ת.ז. {id_number} לא במערכת"}, 404

        cur.execute("""
            UPDATE examinees SET is_present = 1, scan_time = %s, scanner_technician = %s WHERE id_number = %s
        """, (datetime.now(), session.get('username',''), id_number))
        conn.commit()
        return {"success": True, "name": name, "id": id_number}
    except Exception as e:
        return {"error": str(e)}, 500
    finally:
        release_db_connection(conn)

# ── BEACON: GET endpoint לסריקות מהטלפון (עוקף בעיות SSL בכרום) ──────
@app.route('/api/exam-scan-beacon', methods=['GET'])
@login_required
def api_exam_scan_beacon():
    """
    מקבל נתוני סריקה דרך GET params ומחזיר pixel שקוף.
    הדפדפן תמיד שולח GET לתמונות — עוקף בעיות SSL בכרום מובייל.
    """
    qr_text    = (request.args.get('qr', '') or '').strip()
    computer   = (request.args.get('computer', '') or '').strip()
    pc_status  = (request.args.get('pc_status', '') or '').strip()
    is_present = int(request.args.get('is_present', 1) or 1)
    seat       = (request.args.get('seat', '') or '').strip()
    col        = (request.args.get('col', '') or '').strip()
    technician = session.get('username', '')

    parts = qr_text.split('|')
    if len(parts) >= 3:
        id_number  = parts[1] if len(parts) > 1 else ''
        full_name  = parts[2] if len(parts) > 2 else ''
        exam_name  = parts[0] if parts[0] != 'EXAMINEE' else ''
        scan_time_str = datetime.now().strftime('%d/%m/%Y %H:%M:%S')

        def _save():
            try:
                import gspread
                from google.oauth2.service_account import Credentials
                scopes   = ['https://www.googleapis.com/auth/spreadsheets', 'https://www.googleapis.com/auth/drive']
                sa_file  = os.getenv('GOOGLE_SERVICE_ACCOUNT_FILE', 'service_account.json')
                sheet_id = os.getenv('EXAM_ATTENDANCE_SHEET_ID', '1YWLJA5T8Uq7IGzlzXSA1PPwrdSIPx9eazEcwWXwh3uM')
                creds    = Credentials.from_service_account_file(sa_file, scopes=scopes)
                client   = gspread.authorize(creds)
                ws       = client.open_by_key(sheet_id).sheet1
                row = [exam_name, id_number, full_name, computer, pc_status,
                       '✅' if is_present == 1 else '❌', scan_time_str, technician, col, seat]
                ws.append_row(row, value_input_option='USER_ENTERED')
                print(f"[BEACON] Saved: {full_name} | {computer}")
            except Exception as ex:
                print(f"[BEACON] Save error: {ex}")
        threading.Thread(target=_save, daemon=True).start()
    else:
        print(f"[BEACON] ⚠️ Invalid QR: {qr_text[:40]}")

    # החזר pixel שקוף 1x1
    pixel = base64.b64decode('R0lGODlhAQABAAAAACH5BAEKAAEALAAAAAABAAEAAAICTAEAOw==')
    return send_file(io.BytesIO(pixel), mimetype='image/gif', max_age=0)

@app.route('/api/check-computer-used', methods=['POST'])
@login_required
def api_check_computer_used():
    """בדיקה אם מחשב כבר שויך לנבחן אחר בגיליון הנוכחי"""
    global attendance_cache
    data = request.json or {}
    computer = (data.get('computer', '') or '').strip()
    exam_name = (data.get('exam_name', '') or '').strip()

    if not computer or not exam_name:
        return jsonify({"in_use": False})

    # חיפוש מהיר במטמון (אם כבר טעון)
    # אם לא טעון — ניגש לגוגל שיטס
    try:
        import gspread
        from google.oauth2.service_account import Credentials
        conn_db = get_db_connection()
        sheet_id = None
        if conn_db:
            try:
                cur_db = get_safe_cursor(conn_db)
                cur_db.execute("SELECT sheets_id FROM projects WHERE name = %s OR keywords LIKE %s", (exam_name, f"%{exam_name}%"))
                row_proj = cur_db.fetchone()
                if row_proj:
                    sheet_id = row_proj['sheets_id']
                cur_db.close()
            except Exception:
                pass
            finally:
                release_db_connection(conn_db)
        if not sheet_id:
            sheet_id = os.getenv('EXAM_ATTENDANCE_SHEET_ID', '1YWLJA5T8Uq7IGzlzXSA1PPwrdSIPx9eazEcwWXwh3uM')

        scopes = ['https://www.googleapis.com/auth/spreadsheets', 'https://www.googleapis.com/auth/drive']
        sa_file = os.getenv('GOOGLE_SERVICE_ACCOUNT_FILE', 'service_account.json')
        creds = Credentials.from_service_account_file(sa_file, scopes=scopes)
        client = gspread.authorize(creds)
        try:
            ws = client.open_by_key(sheet_id).worksheet(exam_name)
        except Exception:
            ws = client.open_by_key(sheet_id).sheet1
        all_rows = ws.get_all_values()
        for r in all_rows[1:]:
            # עמודה K (אינדקס 10) = מ.מחשב, עמודה L (אינדקס 11) = נוכחות
            if len(r) > 10 and str(r[10]).strip() == str(computer).strip():
                if len(r) > 11 and r[11].strip():
                    existing_name = r[0] if r[0] else 'נבחן'
                    return jsonify({"in_use": True, "name": existing_name})
        return jsonify({"in_use": False})
    except Exception as ex:
        print(f"[check-computer-used] Error: {ex}", flush=True)
        return jsonify({"in_use": False})


@app.route('/api/exam-scan-double', methods=['POST'])
@login_required
def api_exam_scan_double():
    """סריקה כפולה: נבחן + מחשב + סטטוסים — מקבל JSON או form data"""
    global attendance_cache

    if request.is_json:
        data = request.json
    else:
        data = request.form
    qr_text   = (data.get('qr', '') or '').strip()
    computer  = (data.get('computer', '') or '').strip()
    seat      = (data.get('seat', '') or '').strip()
    col       = (data.get('col', '') or '').strip()
    pc_status = (data.get('pc_status', '') or '').strip()
    is_present = int(data.get('is_present', 1) or 1)

    parts = qr_text.split('|')
    if len(parts) < 3:
        return {"error": "QR לא מזוהה כנבחן"}, 400

    # קריאת נתונים ישירות מה-QR - לא תלוי ב-DB
    id_number = parts[1] if len(parts) > 1 else ''
    full_name = parts[2] if len(parts) > 2 else ''
    
    # שם המבחן — מהפרמטר, ובמידת הצורך מה-QR (parts[0])
    exam_name = (data.get('exam_name', '') or '').strip()
    if not exam_name:
        exam_name = parts[0] if (len(parts) > 0 and parts[0] != 'EXAMINEE') else ''

    scan_time_str = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
    technician = session.get('username', '')

    # ── דינמי: איתור קובץ ה-Google Sheet הייעודי למבחן זה ──
    sheet_id = None
    if exam_name:
        conn_db = get_db_connection()
        if conn_db:
            try:
                cur_db = get_safe_cursor(conn_db)
                cur_db.execute("SELECT sheets_id FROM projects WHERE name = %s OR keywords LIKE %s", (exam_name, f"%{exam_name}%"))
                row_proj = cur_db.fetchone()
                if row_proj:
                    sheet_id = row_proj['sheets_id']
                cur_db.close()
            except Exception as ex_db:
                print(f"Error looking up project sheets_id: {ex_db}")
            finally:
                release_db_connection(conn_db)

    # Fallback למפתח ברירת המחדל
    if not sheet_id:
        sheet_id = os.getenv('EXAM_ATTENDANCE_SHEET_ID', '1YWLJA5T8Uq7IGzlzXSA1PPwrdSIPx9eazEcwWXwh3uM')

    # שמירת מפתח הגיליון והבחינה בסשן לטובת ביטול
    session['last_exam_sheet_id'] = sheet_id
    session['last_exam_name'] = exam_name
    session['last_id_number'] = id_number

    # טעינת המטמון במידה ולא נטען עדיין עבור הבחינה הנוכחית
    if exam_name not in attendance_cache:
        print(f"[CACHE] Loading attendance cache for exam: {exam_name}", flush=True)
        try:
            import gspread
            from google.oauth2.service_account import Credentials
            scopes = ['https://www.googleapis.com/auth/spreadsheets','https://www.googleapis.com/auth/drive']
            sa_file  = os.getenv('GOOGLE_SERVICE_ACCOUNT_FILE', 'service_account.json')
            creds    = Credentials.from_service_account_file(sa_file, scopes=scopes)
            client   = gspread.authorize(creds)
            # חיפוש לפי כותרת (שם) קודם, אחרכך על פי מפתח
            ws = None
            if exam_name:
                try:
                    ws = client.open(exam_name).sheet1
                    print(f"[CACHE] Opened spreadsheet by title: {exam_name}", flush=True)
                except Exception:
                    pass
            if ws is None:
                try:
                    ws = client.open_by_key(sheet_id).worksheet(exam_name)
                except Exception:
                    ws = client.open_by_key(sheet_id).sheet1
            all_rows = ws.get_all_values()

            # מבנה עמודות לפי האקסל:
            # [0]שם פרטי | [1]שם משפחה | [2]ת.ז | [3]התאמות | [4]סיסמה | [5]שם משתמש | [6]גרסה | [7]סה' אולם | [8]טור | [9]כסא | [10]מ.מחשב | [11]נוכחות | [12]שעת סריקה | [13]טכנאי
            attendance_cache[exam_name] = {}
            for r in all_rows[1:]:
                # בדיקה בעמודה L (אינדקס 11) = נוכחות
                if len(r) > 11 and r[11].strip() and len(r) > 2:
                    attendance_cache[exam_name][str(r[2]).strip()] = r[12] if (len(r) > 12 and r[12].strip()) else 'נוכח'
            print(f"[CACHE] Loaded {len(attendance_cache[exam_name])} active scans for {exam_name}", flush=True)
        except Exception as ex:
            print(f"[CACHE ERROR] Failed to build cache for {exam_name}: {ex}", flush=True)
            attendance_cache[exam_name] = {}

    # שמירה תמיד — ללא בדיקת כפולים
    attendance_cache.setdefault(exam_name, {})[str(id_number).strip()] = scan_time_str

    def save_to_exam_sheet_in_thread(target_sheet_id, target_exam_name, target_id_number, target_full_name, target_computer, target_col, target_seat, target_pc_status, target_scan_time, target_technician):
        import sys, traceback, gspread
        from google.oauth2.service_account import Credentials
        print(f"[THREAD] Starting Google Sheets update for {target_full_name} ({target_id_number})...", flush=True)
        try:
            scopes = ['https://www.googleapis.com/auth/spreadsheets','https://www.googleapis.com/auth/drive']
            sa_file  = os.getenv('GOOGLE_SERVICE_ACCOUNT_FILE', 'service_account.json')
            creds    = Credentials.from_service_account_file(sa_file, scopes=scopes)
            client   = gspread.authorize(creds)
            # חיפוש לפי כותרת (שם) קודם, אחרכך על פי מפתח
            ws = None
            if target_exam_name:
                try:
                    ws = client.open(target_exam_name).sheet1
                    print(f"[THREAD] Opened by title: {target_exam_name}", flush=True)
                except Exception:
                    pass
            if ws is None:
                try:
                    ws = client.open_by_key(target_sheet_id).worksheet(target_exam_name)
                except Exception:
                    ws = client.open_by_key(target_sheet_id).sheet1

            # תמיד מוסיף שורה חדשה — כל סריקה = שורה חדשה
            new_row = [
                target_full_name,    # [0] שם פרטי (שם מלא)
                "",                  # [1] שם משפחה
                target_id_number,    # [2] ת.ז
                "",                  # [3] התאמות
                "",                  # [4] סיסמה
                "",                  # [5] שם משתמש
                target_exam_name,    # [6] גרסה/בחינה
                "",                  # [7] סה' אולם
                target_col,          # [8] טור
                target_seat,         # [9] כסא
                target_computer,     # [10] מ.מחשב ← נסרק
                "1",                 # [11] נוכחות ← נסרק
                target_scan_time,    # [12] שעת סריקה
                target_technician    # [13] טכנאי
            ]
            ws.append_row(new_row, value_input_option='USER_ENTERED')
            print(f"[THREAD OK] Appended new row for {target_full_name}", flush=True)

        except Exception as ex:
            print(f"[THREAD ERROR] Save failed for {target_full_name}: {ex}", flush=True)
            traceback.print_exc()

    import threading
    t = threading.Thread(
        target=save_to_exam_sheet_in_thread,
        args=(sheet_id, exam_name, id_number, full_name, computer, col, seat, pc_status, scan_time_str, technician),
        daemon=False
    )
    t.start()

    return jsonify({"success": True})



@app.route('/api/undo-last-scan', methods=['POST'])
@login_required
def undo_last_scan():
    """ביטול סריקה אחרונה וניקוי סטטוס נוכחות בגוגל שיטס"""
    global attendance_cache
    try:
        import gspread, os
        from google.oauth2.service_account import Credentials
        scopes   = ['https://www.googleapis.com/auth/spreadsheets',
                    'https://www.googleapis.com/auth/drive']
        sa_file  = os.getenv('GOOGLE_SERVICE_ACCOUNT_FILE', 'service_account.json')
        creds    = Credentials.from_service_account_file(sa_file, scopes=scopes)
        client   = gspread.authorize(creds)
        
        sheet_id = session.get('last_exam_sheet_id') or os.getenv('EXAM_ATTENDANCE_SHEET_ID', '1YWLJA5T8Uq7IGzlzXSA1PPwrdSIPx9eazEcwWXwh3uM')
        exam_name = session.get('last_exam_name')
        last_id_number = session.get('last_id_number')
        
        # הסרה מהמטמון
        if exam_name and last_id_number and exam_name in attendance_cache:
            if last_id_number in attendance_cache[exam_name]:
                del attendance_cache[exam_name][last_id_number]
                print(f"[UNDO CACHE] Removed {last_id_number} from cache for {exam_name}", flush=True)

        try:
            ws = client.open_by_key(sheet_id).worksheet(exam_name)
        except Exception:
            ws = client.open_by_key(sheet_id).sheet1

        if last_id_number:
            all_vals = ws.get_all_values()
            row_idx = -1
            for idx, r in enumerate(all_vals):
                if len(r) > 1 and str(r[1]).strip() == str(last_id_number).strip():
                    row_idx = idx + 1
                    break
            
            if row_idx != -1:
                r_new = list(all_vals[row_idx - 1])
                while len(r_new) < 13:
                    r_new.append("")
                # איפוס עמודות: מחשב, נוכחות, מצב מחשב, שעת סריקה
                r_new[6]  = ""
                r_new[10] = ""
                r_new[11] = ""
                r_new[12] = ""
                ws.update(f"A{row_idx}", [r_new], value_input_option='USER_ENTERED')
                print(f"[UNDO OK] Cleared presence on row {row_idx} for ID {last_id_number}", flush=True)
                return {"success": True}
        
        # Fallback לביטול השורה האחרונה במידה ולא נמצא נתון ספציפי
        all_vals = ws.get_all_values()
        last_row = len(all_vals)
        if last_row > 1:
            ws.delete_rows(last_row)
            return {"success": True}
            
        return {"success": False, "error": "לא נמצא נבחן לביטול"}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.route('/api/exam-scan', methods=['POST'])
@login_required
def api_exam_scan():
    """סריקת QR לנוכחות"""
    data = request.json
    qr_text = data.get('qr', '').strip()
    if not qr_text:
        return {"error": "לא התקבל QR"}, 400

    # פרמט: שם_בחינה|ת.ז.|שם|קוד|סיסמא|טור|כסא
    parts = qr_text.split('|')
    if len(parts) < 6:
        return {"error": "QR לא מזוהה כנבחן", "type": "unknown"}, 400

    id_number = parts[1] if len(parts) > 1 else ''

    conn = get_db_connection()
    if not conn: return {"error": "DB Error"}, 500
    try:
        cur = get_safe_cursor(conn)
        cur.execute("SELECT * FROM examinees WHERE id_number = %s", (id_number,))
        examinee = cur.fetchone()
        if not examinee:
            return {"error": f"נבחן עם ת.ז. {id_number} לא נמצא במערכת"}, 404

        examinee = dict(examinee)
        if examinee.get('is_present') in (True, 1):
            return {"success": True, "already": True, "examinee": examinee}

        # סמן כנוכח
        cur.execute("""
            UPDATE examinees SET is_present = 1, scan_time = %s, scanner_technician = %s WHERE id_number = %s
        """, (datetime.now(), session.get('username',''), id_number))
        conn.commit()

        # סנכרון לגוגל דרייב ברקע
        exam_name = examinee.get('exam_name')
        if exam_name:
            cur.execute("SELECT * FROM examinees WHERE exam_name = %s", (exam_name,))
            all_exam_examinees = [dict(row) for row in cur.fetchall()]
            
            import threading
            from sync_attendance_drive import sync_exam_to_drive
            threading.Thread(target=sync_exam_to_drive, args=(exam_name, all_exam_examinees)).start()

        cur.close()

        # סנכרון Google Sheets (קוד קיים)
        threading.Thread(target=sync_inventory_to_sheets, daemon=True).start()

        examinee['is_present'] = True
        examinee['attend_time'] = datetime.now().strftime("%H:%M:%S")
        return {"success": True, "already": False, "examinee": examinee}
    except Exception as e:
        return {"error": str(e)}, 500
    finally:
        release_db_connection(conn)

@app.route('/exam-attendance/delete/<int:eid>', methods=['POST'])
@login_required
def exam_attendance_delete(eid):
    """מחיקת נבחן"""
    conn = get_db_connection()
    if not conn: return "DB Error", 500
    try:
        cur = get_safe_cursor(conn)
        cur.execute("DELETE FROM examinees WHERE id = %s", (eid,))
        conn.commit()
        cur.close()
        flash("הנבחן נמחק בהצלחה", "success")
    finally:
        release_db_connection(conn)
    return redirect(url_for('exam_attendance'))

@app.route('/exam-attendance/clear', methods=['POST'])
@login_required
def exam_attendance_clear():
    """איפוס כל הנוכחות (לפני בחינה חדשה) - מנהל בלבד"""
    if session.get('role') != 'admin':
        flash("אין הרשאה לפעולה זו", "danger")
        return redirect(url_for('exam_attendance'))
    conn = get_db_connection()
    if not conn: return "DB Error", 500
    try:
        cur = get_safe_cursor(conn)
        cur.execute("UPDATE examinees SET is_present = 0, scan_time = NULL")
        conn.commit()
        cur.close()
        flash("✅ כל הנוכחות אופסה – מוכן לבחינה חדשה!", "success")
    finally:
        release_db_connection(conn)
    return redirect(url_for('exam_attendance'))

@app.route('/exam-attendance/scanner')
@login_required
def exam_attendance_scanner():
    """עמוד סריקת נוכחות"""
    return render_template('exam_scanner.html')


@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

# --- Standalone Exam Form Generator ---

@app.route('/exam-generator')
def exam_generator():
    """דף העלאת אקסל ליצירת טפסי בחינה"""
    return render_template('exam_generator.html')

@app.route('/generate-forms', methods=['POST'])
def generate_forms():
    """מעבד אקסל ומחזיר דף מוכן להדפסה עם QR"""
    excel_file = request.files.get('excel_file')
    if not excel_file:
        return "לא הועלה קובץ", 400

    try:
        wb = openpyxl.load_workbook(excel_file)
        ws = wb.active

        headers = [str(cell.value).strip() if cell.value else '' for cell in ws[1]]

        col_map = {}
        for i, h in enumerate(headers):
            h_l = h.lower()
            if 'שם' in h and 'נבחן' in h: col_map['name'] = i
            elif 'שם' in h and col_map.get('name') is None: col_map['name'] = i
            elif 'זהות' in h or 'ת.ז' in h or 'id' in h_l: col_map['id_number'] = i
            elif 'משתמש' in h or 'קוד' in h or 'user' in h_l: col_map['username'] = i
            elif 'סיסמ' in h or 'pass' in h_l: col_map['password'] = i
            elif 'התאמ' in h or 'notes' in h_l: col_map['notes'] = i
            elif 'מחשב' in h or 'computer' in h_l or 'מספר' in h: col_map['computer'] = i
            elif 'מיקום' in h or 'כיתה' in h or 'location' in h_l: col_map['location'] = i

        students = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            if not any(row): continue
            def get_val(key):
                idx = col_map.get(key)
                if idx is None: return ''
                v = row[idx]
                return str(v).strip() if v is not None else ''

            name = get_val('name')
            if not name or name == 'None': continue

            students.append({
                'name': name,
                'id_number': get_val('id_number'),
                'username': get_val('username'),
                'password': get_val('password'),
                'notes': get_val('notes'),
                'computer': get_val('computer'),
                'location': get_val('location'),
            })

        return render_template('exam_forms_print.html', students=students)

    except Exception as e:
        import traceback; traceback.print_exc()
        return f"שגיאה: {e}", 500

# ── FAULT REPORT: טופס תקלות מחשב ────────────────────────────────
@app.route('/fault-report', methods=['GET'])
@login_required
def fault_report_page():
    """עמוד טופס דיווח תקלה"""
    # שלוף רשימת מחשבים לרשימת auto-complete
    conn = get_db_connection()
    barcodes = []
    if conn:
        try:
            cur = get_safe_cursor(conn)
            cur.execute("SELECT barcode, case_number, location FROM computers ORDER BY barcode")
            barcodes = [dict(r) for r in cur.fetchall()]
            cur.close()
        except Exception:
            pass
        finally:
            release_db_connection(conn)
    return render_template('fault_report.html', barcodes=barcodes)

@app.route('/api/submit-fault', methods=['POST'])
@login_required
def api_submit_fault():
    """קבלת דיווח תקלה — שומר לשיטס + להיסטוריה"""
    data = request.json or {}
    barcode     = (data.get('barcode', '') or '').strip()
    fault_type  = (data.get('fault_type', '') or '').strip()
    description = (data.get('description', '') or '').strip()
    location    = (data.get('location', '') or '').strip()
    technician  = session.get('username', '')
    report_time = datetime.now().strftime('%d/%m/%Y %H:%M:%S')

    if not barcode:
        return jsonify({"success": False, "error": "חובה להזין מספר מחשב"}), 400
    if not description:
        return jsonify({"success": False, "error": "חובה לתאר את התקלה"}), 400

    # ── שמירה לגיליון 'תקלות' בשיטס ─────────────────────────────
    def save_fault_to_sheets():
        import traceback
        print(f"[FAULT] ▶️ שומר תקלה: מחשב {barcode} | {fault_type}", flush=True)
        try:
            import gspread
            from google.oauth2.service_account import Credentials
            scopes   = ['https://www.googleapis.com/auth/spreadsheets', 'https://www.googleapis.com/auth/drive']
            sa_file  = os.getenv('GOOGLE_SERVICE_ACCOUNT_FILE', 'service_account.json')
            sheet_id = os.getenv('GOOGLE_SHEETS_ID')
            if not sheet_id:
                print("[FAULT] ⚠️ GOOGLE_SHEETS_ID לא מוגדר", flush=True)
                return
            creds  = Credentials.from_service_account_file(sa_file, scopes=scopes)
            client = gspread.authorize(creds)
            sh     = client.open_by_key(sheet_id)
            # פתח/צור גיליון 'תקלות'
            try:
                ws = sh.worksheet('תקלות')
            except gspread.WorksheetNotFound:
                ws = sh.add_worksheet(title='תקלות', rows='500', cols='8')
                ws.append_row(['תאריך', 'מחשב', 'סוג תקלה', 'תיאור', 'מיקום', 'טכנאי', 'סטטוס טיפול'],
                              value_input_option='USER_ENTERED')
                ws.format('A1:G1', {'textFormat': {'bold': True},
                                    'backgroundColor': {'red': 1.0, 'green': 0.85, 'blue': 0.4}})
            ws.append_row([report_time, barcode, fault_type, description, location, technician, 'ממתין לטיפול'],
                          value_input_option='USER_ENTERED')
            print(f"[FAULT] ✅ נשמר לגיליון תקלות: {barcode}", flush=True)
        except Exception as ex:
            print(f"[FAULT] ❌ שגיאה בשמירה לשיטס: {ex}", flush=True)
            traceback.print_exc()

    threading.Thread(target=save_fault_to_sheets, daemon=False).start()

    # ── שמירה להיסטוריה בDB ───────────────────────────────────────
    conn = get_db_connection()
    if conn:
        try:
            cur = get_safe_cursor(conn)
            cur.execute("SELECT id FROM computers WHERE barcode = %s", (barcode,))
            comp = cur.fetchone()
            if comp:
                comp_id = comp['id']
                note_text = f"[תקלה] {fault_type}: {description}"
                cur.execute("""
                    INSERT INTO inventory_history (computer_id, technician, change_type, old_value, new_value)
                    VALUES (%s, %s, 'דיווח תקלה', %s, %s)
                """, (comp_id, technician,
                       f"מיקום: {location}",
                       note_text))
                conn.commit()
            cur.close()
        except Exception as e:
            print(f"[FAULT] DB history error: {e}", flush=True)
        finally:
            release_db_connection(conn)

    return jsonify({"success": True, "message": f"תקלה דווחה בהצלחה עבור מחשב {barcode}"})

@app.route('/api/sheets-sync-status')
@login_required
def api_sheets_sync_status():
    """מחזיר זמן הסנכרון האחרון מגיליון שיטס"""
    global _last_sheets_import
    if _last_sheets_import:
        diff = (datetime.now() - _last_sheets_import).seconds
        if diff < 60:
            ago = f"לפני {diff} שניות"
        else:
            ago = f"לפני {diff // 60} דקות"
        return jsonify({"last_sync": _last_sheets_import.strftime('%H:%M:%S'), "ago": ago})
    return jsonify({"last_sync": None, "ago": "טרם סונכרן"})

# ── CAGE INFO PAGE (FOR MOBILE/PHONE QR SCAN) ─────────────────────────
@app.route('/cage-info/<cage_id>')
@login_required
def cage_info_page(cage_id):
    conn = get_db_connection()
    if not conn: return "DB Error", 500
    try:
        cur = get_safe_cursor(conn)
        
        # Get cage details
        cur.execute("SELECT * FROM cages WHERE cage_id = %s", (cage_id,))
        cage = cur.fetchone()
        
        # Get computers in cage
        cur.execute("""
            SELECT id, barcode, case_number, status, location, specs, scan_time, notes
            FROM computers
            WHERE cage_number = %s OR cage_name = %s
            ORDER BY scan_time DESC NULLS LAST
        """, (cage_id, cage_id))
        computers = cur.fetchall()
        
        cur.close()
        
        if not cage:
            cage = {
                'cage_id': cage_id,
                'name': f'כלוב {cage_id}',
                'location': '',
                'notes': 'כלוב זה נוצר אוטומטית בעת סריקת מחשבים.'
            }
            
        return render_template('cage_info.html', cage=cage, computers=computers, total=len(computers))
    except Exception as e:
        print(f"Error in cage_info_page: {e}")
        return f"<h1>Error: {e}</h1>", 500
    finally:
        release_db_connection(conn)


@app.route('/cages/print/<cage_id>')
@login_required
def print_cage_page(cage_id):
    conn = get_db_connection()
    if not conn: return "DB Error", 500
    try:
        cur = get_safe_cursor(conn)
        
        # Get cage details
        cur.execute("SELECT * FROM cages WHERE cage_id = %s", (cage_id,))
        cage = cur.fetchone()
        
        # Get computers in cage
        cur.execute("""
            SELECT barcode, case_number, status, location
            FROM computers
            WHERE cage_number = %s OR cage_name = %s
            ORDER BY scan_time DESC NULLS LAST
        """, (cage_id, cage_id))
        computers = cur.fetchall()
        
        cur.close()
        
        if not cage:
            cage = {
                'cage_id': cage_id,
                'name': f'כלוב {cage_id}',
                'location': '',
                'notes': ''
            }
            
        return render_template('print_cage.html', cage=cage, computers=computers, total=len(computers))
    except Exception as e:
        print(f"Error in print_cage_page: {e}")
        return f"<h1>Error: {e}</h1>", 500
    finally:
        release_db_connection(conn)

# --- End of Routes ---

if __name__ == '__main__':
    import socket
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        local_ip = s.getsockname()[0]
        s.close()
    except Exception:
        local_ip = "127.0.0.1"

    print("\n[START] URI SYSTEM IS LIVE! (HTTPS ENABLED FOR MOBILE SCANNER)")
    print("URL Link: https://127.0.0.1:5000")
    print(f"Mobile Link: https://{local_ip}:5000\n")
    print("[WARNING] When opening on iPhone, you will see a 'Not Private' warning. Click 'Show Details' -> 'Visit this website' to bypass and test the scanner.")
    import os
    cert_file = os.path.join(os.path.dirname(__file__), 'server.crt')
    key_file  = os.path.join(os.path.dirname(__file__), 'server.key')
    if os.path.exists(cert_file) and os.path.exists(key_file):
        ssl_ctx = (cert_file, key_file)
        print("[OK] Using custom SSL certificate (valid for local network)")
    else:
        ssl_ctx = 'adhoc'
        print("[WARN] Custom cert not found, using adhoc SSL")
    app.run(host='0.0.0.0', debug=True, port=5000, ssl_context=ssl_ctx)
